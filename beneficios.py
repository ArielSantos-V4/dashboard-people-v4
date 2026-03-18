import streamlit as st
import pandas as pd
import altair as alt
import re
import unicodedata
from datetime import datetime, date
from docx import Document
import gspread
from google.oauth2.service_account import Credentials
import os

# ==========================================
# FUNÇÕES AUXILIARES
# ==========================================
MESES_PT = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
    5: "maio", 6: "junho", 7: "julho", 8: "agosto",
    9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

def substituir_texto(paragraphs, mapa):
    for p in paragraphs:
        for run in p.runs:
            for chave, valor in mapa.items():
                if chave in run.text:
                    run.text = run.text.replace(chave, str(valor))

def formatar_cnpj(valor):
    if pd.isna(valor) or valor == "":
        return ""
    v = str(valor).replace(".0", "").replace(".", "").replace("-", "").replace("/", "").strip()
    v = v.zfill(14)
    if len(v) == 14:
        return f"{v[:2]}.{v[2:5]}.{v[5:8]}/{v[8:12]}-{v[12:]}"
    return v

def normalizar_cpf(valor):
    if pd.isna(valor) or valor == "":
        return ""
    v = str(valor).replace(".0", "").replace(".", "").replace("-", "").replace("/", "").strip()
    return re.sub(r"\D", "", v).zfill(11)

def email_para_nome_arquivo(email):
    if not email:
        return ""
    return str(email).replace("@", "_").replace(".", "_").lower()

def carregar_desligados_google_sheets():
    # Tenta carregar credenciais usando st.secrets (mais seguro e correto para o Streamlit Cloud)
    try:
        scopes = [
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive"
        ]
        
        # Aqui usamos st.secrets em vez de procurar o arquivo .json
        creds = Credentials.from_service_account_info(
            st.secrets["gcp_service_account"], 
            scopes=scopes
        )
        
        client = gspread.authorize(creds)
        
        # ID da sua planilha Master
        spreadsheet = client.open_by_key("13EPwhiXgh8BkbhyrEy2aCy3cv1O8npxJ_hA-HmLZ-pY") 
        
        # Acessa a aba de desligados pelo GID (1422602176)
        # O gspread não tem "get_worksheet_by_id" nativo, então fazemos esse loop rápido:
        worksheet = None
        for sheet in spreadsheet.worksheets():
            if str(sheet.id) == "1422602176":
                worksheet = sheet
                break
        
        if worksheet:
            dados = worksheet.get_all_records()
            return pd.DataFrame(dados)
        else:
            st.error("Aba de desligados (GID 1422602176) não encontrada.")
            return pd.DataFrame()
            
    except Exception as e:
        st.error(f"Erro ao conectar com o Google Sheets: {e}")
        return pd.DataFrame()

# ==========================================
# MODAIS (GLOBAL)
# ==========================================

@st.dialog("📄 Gerar Inclusão Subfatura")
def modal_inclusao_subfatura(df):
    nomes = sorted(df["Nome"].dropna().unique())
    nome_escolhido = st.selectbox("Selecione o investidor", nomes, key="nome_subfatura")
    data_vigencia = st.date_input("Data de início da vigência", format="DD/MM/YYYY")

    st.markdown("<br>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    
    if col2.button("✅ Gerar", use_container_width=True, key="btn_subfatura"):
        dados = df[df["Nome"] == nome_escolhido].iloc[0]
        
        razao_social = str(dados.get("Razão social", ""))
        cnpj = formatar_cnpj(dados.get("CNPJ", ""))
        cpf = normalizar_cpf(dados.get("CPF", ""))
        email_pessoal = str(dados.get("E-mail pessoal", ""))
        email_arquivo = email_para_nome_arquivo(email_pessoal)
        modelo_contrato = str(dados.get("Modelo de contrato", ""))

        if "PJ" not in modelo_contrato.upper():
            st.warning(f"⚠️ **{nome_escolhido}** não possui contrato PJ. Modelo atual: **{modelo_contrato}**")

        try:
            doc = Document("Subfatura.docx")
            vigencia_formatada = data_vigencia.strftime("%d/%m/%Y")
            hoje = date.today()
            data_assinatura = f"{hoje.day} de {MESES_PT[hoje.month]} de {hoje.year}"

            mapa = {
                "{RAZAO_SOCIAL}": razao_social,
                "{CNPJ}": cnpj,
                "{VIGENCIA}": vigencia_formatada,
                "{DATA}": data_assinatura
            }

            substituir_texto(doc.paragraphs, mapa)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        substituir_texto(cell.paragraphs, mapa)
            for section in doc.sections:
                substituir_texto(section.header.paragraphs, mapa)

            cpf_limpo = re.sub(r"\D", "", cpf)
            nome_arquivo = f"{nome_escolhido} __ {cpf_limpo} __ {email_arquivo} __ Inclusão Subfatura.docx"
            doc.save(nome_arquivo)

            with open(nome_arquivo, "rb") as f:
                st.download_button("⬇️ Download", f, file_name=nome_arquivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            
            st.link_button("🔁 Converter PDF", "https://www.ilovepdf.com/pt/word_para_pdf", use_container_width=True)
            st.success("Inclusão Subfatura gerada com sucesso ✅")

        except Exception as e:
            st.error(f"Erro ao gerar documento: {e}")

@st.dialog("📄 Gerar Termo de Subestipulante")
def modal_subestipulante(df):
    nomes = sorted(df["Nome"].dropna().unique())
    nome_escolhido = st.selectbox("Selecione o investidor", nomes, key="nome_termo_sub")

    col1, col2, col3 = st.columns([1, 2, 1])
    if col2.button("✅ Gerar Termo", use_container_width=True, key="btn_termo_sub"):
        dados = df[df["Nome"] == nome_escolhido].iloc[0]
        razao_social = str(dados.get("Razão social", ""))
        cnpj = formatar_cnpj(dados.get("CNPJ", ""))
        cpf = normalizar_cpf(dados.get("CPF", ""))
        email_pessoal = str(dados.get("E-mail pessoal", ""))
        email_arquivo = email_para_nome_arquivo(email_pessoal)

        try:
            doc = Document("Termo de integração de subestipulante.docx")
            hoje = date.today()
            data_assinatura = f"{hoje.day} de {MESES_PT[hoje.month]} de {hoje.year}"

            mapa = {"{RAZAO_SOCIAL}": razao_social, "{CNPJ}": cnpj, "{DATA}": data_assinatura}

            substituir_texto(doc.paragraphs, mapa)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        substituir_texto(cell.paragraphs, mapa)
            for section in doc.sections:
                substituir_texto(section.header.paragraphs, mapa)

            cpf_limpo = re.sub(r"\D", "", cpf)
            nome_arquivo = f"{nome_escolhido} __ {cpf_limpo} __ {email_arquivo} __ Termo Subestipulante.docx"
            doc.save(nome_arquivo)

            with open(nome_arquivo, "rb") as f:
                st.download_button("⬇️ Download", f, file_name=nome_arquivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            
            st.link_button("🔁 Converter PDF", "https://www.ilovepdf.com/pt/word_para_pdf", use_container_width=True)
            st.success("Termo de Subestipulante gerado com sucesso ✅")
        except Exception as e:
            st.error(f"Erro ao gerar documento: {e}")

@st.dialog("📄 Gerar Termo de Não Adesão")
def modal_nao_adesao(df):
    nomes = sorted(df["Nome"].dropna().unique())
    nome_escolhido = st.selectbox("Selecione o investidor", nomes, key="nome_nao_adesao")

    col1, col2, col3 = st.columns([1, 2, 1])
    if col2.button("✅ Gerar Termo", use_container_width=True, key="btn_nao_adesao"):
        dados = df[df["Nome"] == nome_escolhido].iloc[0]
        razao_social = str(dados.get("Razão social", ""))
        cnpj = formatar_cnpj(dados.get("CNPJ", ""))
        
        try:
            doc = Document("Termo de não adesão - Plano de Saúde e Dental.docx")
            hoje = date.today()
            data_assinatura = f"{hoje.day} de {MESES_PT[hoje.month]} de {hoje.year}"
            
            mapa = {"{RAZAO_SOCIAL}": razao_social, "{CNPJ}": cnpj, "{DATA}": data_assinatura}

            substituir_texto(doc.paragraphs, mapa)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        substituir_texto(cell.paragraphs, mapa)
            for section in doc.sections:
                substituir_texto(section.header.paragraphs, mapa)
                substituir_texto(section.footer.paragraphs, mapa)

            nome_arquivo = f"Termo de não adesão ao plano - {nome_escolhido}.docx"
            doc.save(nome_arquivo)

            with open(nome_arquivo, "rb") as f:
                st.download_button("⬇️ Download", f, file_name=nome_arquivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            
            st.link_button("🔁 Converter PDF", "https://www.ilovepdf.com/pt/word_para_pdf", use_container_width=True)
            st.success("Termo de Não Adesão gerado com sucesso ✅")
        except Exception as e:
            st.error(f"Erro ao gerar documento: {e}")

@st.dialog("📄 Gerar Exclusão Subfatura")
def modal_exclusao_subfatura():
    # Carrega planilha de desligados (função específica)
    df_desligados = carregar_desligados_google_sheets()
    
    if df_desligados.empty:
        st.warning("Não foi possível carregar a base de desligados.")
        return

    nomes = sorted(df_desligados["Nome"].dropna().unique())
    nome_escolhido = st.selectbox("Selecione o investidor", nomes, key="nome_exclusao")
    data_exclusao = st.date_input("Data de exclusão", format="DD/MM/YYYY")

    st.markdown("<br>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    
    if col2.button("✅ Gerar", use_container_width=True, key="btn_exclusao"):
        dados = df_desligados[df_desligados["Nome"] == nome_escolhido].iloc[0]
        
        # 1. Preparação dos dados
        razao_social = str(dados.get("Razão social", "")).upper()
        cnpj = formatar_cnpj(dados.get("CNPJ", ""))
        cpf = normalizar_cpf(dados.get("CPF", ""))
        email_pessoal = str(dados.get("E-mail pessoal", ""))
        email_arquivo = email_para_nome_arquivo(email_pessoal)
        
        hoje = date.today()
        data_assinatura = f"{hoje.day} de {MESES_PT[hoje.month]} de {hoje.year}"

        # 2. Mapa com as chaves EXATAS que você pediu
        mapa = {
            "{{razao_social}}": razao_social,
            "{{cnpj}}": cnpj,
            "{{data_exclusao}}": data_exclusao.strftime("%d/%m/%Y"),
            "{{data}}": data_assinatura
        }

        try:
            # Carrega o modelo
            doc = Document("Exclusão Subfatura.docx")

            # 3. Executa a substituição em parágrafos, tabelas e cabeçalhos
            for p in doc.paragraphs:
                for chave, valor in mapa.items():
                    if chave in p.text:
                        # Substitui no parágrafo inteiro para evitar quebras do Word
                        p.text = p.text.replace(chave, valor)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for chave, valor in mapa.items():
                                if chave in p.text:
                                    p.text = p.text.replace(chave, valor)

            # 4. Salva e disponibiliza para download
            cpf_limpo = re.sub(r"\D", "", cpf)
            nome_arquivo = f"{nome_escolhido} __ {cpf_limpo} __ {email_arquivo} __ Exclusão Subfatura.docx"
            doc.save(nome_arquivo)

            with open(nome_arquivo, "rb") as f:
                st.download_button("⬇️ Download Documento", f, file_name=nome_arquivo, use_container_width=True, type="primary")
            
            st.success("Documento gerado com sucesso! ✅")
            
        except Exception as e:
            st.error(f"Erro ao processar o Word: {e}")

# ==========================================
# FUNÇÃO PRINCIPAL (RENDER)
# ==========================================
def render(df):
    
    # Proteção de Login
    if "authenticated" not in st.session_state or not st.session_state.authenticated:
        st.warning("Você precisa fazer login para acessar esta página.")
        st.stop()

    # CABEÇALHO (PADRÃO V4)
    c_logo, c_texto = st.columns([0.5, 6]) 
    with c_logo:
        st.image("LOGO VERMELHO.png", width=100) 
    with c_texto:
        st.markdown("""
            <div style="display: flex; flex-direction: column; justify-content: center; height: 100px;">
                <h1 style="margin: 0; padding: 0; font-size: 2.2rem; line-height: 1.1;">Gestão de Benefícios</h1>
                <span style="color: grey; font-size: 1.1rem; margin-top: 2px;">V4 Company</span>
            </div>
        """, unsafe_allow_html=True)
    
    # --- CRIAÇÃO DAS 4 ABAS (Nomes corrigidos para evitar NameError) ---
    aba_dashboard, aba_carteirinhas, aba_analytics, aba_acoes = st.tabs([
        "Dashboard", 
        "Carteirinhas", 
        "Analytics", 
        "Ações"
    ])
    
    # ----------------------------------------------------
    # 1. ABA DASHBOARD
    # ----------------------------------------------------
    with aba_dashboard:
        st.markdown("""
            <div style="background-color: #f1f3f5; padding: 12px; border-radius: 6px; border-left: 5px solid #404040; margin-bottom: 20px;">
                <span style="color: #404040; font-size: 14px;">
                    Acompanhe abaixo os principais indicadores (KPIs) e gráficos demográficos referentes exclusivamente ao <b>plano médico e dental.</b>
                </span>
            </div>
        """, unsafe_allow_html=True)
        
        if "Situação no plano" in df.columns:
            # --- CÁLCULOS DOS KPIs ---
            total_investidores = len(df)
            total_vidas = len(df[df["Situação no plano"] == "Ativo"])
            pendencias = len(df[df["Situação no plano"].isin(["Pendente", "Aguardando docs", "Enviar à DBL"])])
            em_processo = len(df[df["Situação no plano"] == "Aguardando DBL"])
            
            # Novo KPI: Taxa de Adesão
            taxa_adesao = (total_vidas / total_investidores * 100) if total_investidores > 0 else 0
            # Novo KPI: Total Odonto
            total_odonto = len(df[df["Operadora Odonto"].notna() & (df["Operadora Odonto"] != "")])

            # --- PRIMEIRA LINHA DE MÉTRICAS ---
            c1, c2, c3 = st.columns(3)
            c1.metric("Vidas Ativas (Saúde)", total_vidas)
            c2.metric("Pendências Totais", pendencias, delta_color="inverse")
            c3.metric("Em ativação (DBL)", em_processo)
            
            # --- SEGUNDA LINHA DE MÉTRICAS (NOVAS) ---
            c4, c5, c6 = st.columns(3)
            c4.metric("Taxa de Adesão Geral", f"{taxa_adesao:.1f}%")
            c5.metric("Vidas Ativas (Odonto)", total_odonto)
            c6.metric("Total na Base", total_investidores)
            
            st.markdown("---")

            # --- LINHA 1 DE GRÁFICOS (OS QUE VOCÊ JÁ TINHA) ---
            col_g1, col_g2 = st.columns(2)
            with col_g1:
                st.subheader("Situação no plano")
                df_plano = df["Situação no plano"].fillna("Não informado").value_counts().reset_index()
                df_plano.columns = ["Situação", "Quantidade"]
                grafico_pizza = alt.Chart(df_plano).mark_arc(innerRadius=80).encode(
                    theta="Quantidade:Q",
                    color=alt.Color("Situação:N", scale=alt.Scale(range=["#2E8B57", "#FFA500", "#8A2BE2", "#DC143C", "#8B4513", "#808080"])),
                    tooltip=["Situação", "Quantidade"]
                )
                st.altair_chart(grafico_pizza, use_container_width=True)

            with col_g2:
                st.subheader("Vidas por Operadora")
                if "Operadora Médico" in df.columns:
                    df_oper = df[df["Operadora Médico"].notna() & (df["Operadora Médico"] != "")]
                    df_oper_count = df_oper["Operadora Médico"].value_counts().reset_index()
                    df_oper_count.columns = ["Operadora", "Quantidade"]
                    grafico_barras = alt.Chart(df_oper_count).mark_bar(color="#E30613").encode(
                        x=alt.X("Operadora:N", sort="-y"), y="Quantidade:Q"
                    )
                    st.altair_chart(grafico_barras, use_container_width=True)

            # --- LINHA 2 DE GRÁFICOS (NOVOS) ---
            st.markdown("<br>", unsafe_allow_html=True)
            col_g3, col_g4 = st.columns(2)
            
            with col_g3:
                st.subheader("Adesão por Área")
                if "Área" in df.columns:
                    # Filtra apenas ativos para ver quem realmente usa por área
                    df_area = df[df["Situação no plano"] == "Ativo"]["Área"].value_counts().head(10).reset_index()
                    df_area.columns = ["Área", "Vidas"]
                    grafico_area = alt.Chart(df_area).mark_bar(color="#404040").encode(
                        x=alt.X("Vidas:Q"),
                        y=alt.Y("Área:N", sort="-x"),
                        tooltip=["Área", "Vidas"]
                    )
                    st.altair_chart(grafico_area, use_container_width=True)

            with col_g4:
                st.subheader("Adesão por Modelo de Contrato")
                if "Modelo de contrato" in df.columns:
                    df_mod = df[df["Situação no plano"] == "Ativo"]["Modelo de contrato"].value_counts().reset_index()
                    df_mod.columns = ["Modelo", "Vidas"]
                    grafico_modelo = alt.Chart(df_mod).mark_bar(color="#8B0000").encode(
                        x=alt.X("Modelo:N", sort="-y"),
                        y=alt.Y("Vidas:Q"),
                        tooltip=["Modelo", "Vidas"]
                    )
                    st.altair_chart(grafico_modelo, use_container_width=True)

    # ----------------------------------------------------
    # 2. ABA CARTEIRINHAS
    # ----------------------------------------------------
    with aba_carteirinhas:
        st.markdown("""
            <div style="background-color: #f1f3f5; padding: 12px; border-radius: 6px; border-left: 5px solid #404040; margin-bottom: 20px;">
                <span style="color: #404040; font-size: 14px;">
                    Realize consultas de carteirinhas de maneira rápida.
                </span>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown("### 🔎 Consulta de Carteirinhas")
        nome_ben = st.selectbox("Buscar investidor", [""] + sorted(df["Nome"].dropna().unique()), key="sel_ben_cart_v4")
        
        if nome_ben:
            dados = df[df["Nome"] == nome_ben].iloc[0]
            with st.container(border=True):
                c1, c2 = st.columns(2)
                c1.markdown(f"**🏥 Saúde ({dados.get('Operadora Médico', 'N/A')})**")
                c1.code(str(dados.get("Carteirinha médico", "Não possui")).replace(".0", ""), language=None)
                c2.markdown(f"**🦷 Odonto ({dados.get('Operadora Odonto', 'N/A')})**")
                c2.code(str(dados.get("Carteirinha odonto", "Não possui")).replace(".0", ""), language=None)
        
        st.markdown("---")
        
        # --- TABELA DE ATIVOS ---
        st.markdown("### 📋 Base Ativa (Planos de Saúde/Dental)")
        if "Situação no plano" in df.columns:
            # Filtra apenas quem está Ativo
            df_ativos = df[df["Situação no plano"] == "Ativo"].copy()
            
            if not df_ativos.empty:
                # Seleciona colunas relevantes
                colunas_view = ["Nome", "E-mail corporativo"]
                if "Carteirinha médico" in df.columns: colunas_view.append("Carteirinha médico")
                if "Operadora Médico" in df.columns: colunas_view.append("Operadora Médico")
                if "Carteirinha odonto" in df.columns: colunas_view.append("Carteirinha odonto")
                if "Operadora Odonto" in df.columns: colunas_view.append("Operadora Odonto")
                
                # Formata para tirar .0 dos números
                for col in ["Carteirinha médico", "Carteirinha odonto"]:
                    if col in df_ativos.columns:
                        df_ativos[col] = df_ativos[col].astype(str).replace(r'\.0$', '', regex=True)

                st.dataframe(df_ativos[colunas_view], use_container_width=True, hide_index=True)
            else:
                st.info("Nenhum investidor com status 'Ativo' encontrado.")

    # ----------------------------------------------------
    # 3. ABA ANALYTICS
    # ----------------------------------------------------
    with aba_analytics:
        st.markdown("""
            <div style="background-color: #f1f3f5; padding: 12px; border-radius: 6px; border-left: 5px solid #404040; margin-bottom: 20px;">
                <span style="color: #404040; font-size: 14px;">Utilize as abas abaixo para extrair relatórios e acompanhar indicadores os processos de inclusão.</span>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown("### Relatórios Operacionais")
        tabs_rel = st.tabs(["⏰ Pendentes", "📂 Aguardando docs", "📩 Enviar para DBL", "🆗 Ativação"])
        
        with tabs_rel[0]:
            df_p = df[(df["Situação no plano"] == "Pendente") & (df["Modalidade PJ"] != "MEI")]
            st.dataframe(df_p[["Nome", "E-mail corporativo", "Modelo de contrato", "Solicitar documentação"]], use_container_width=True, hide_index=True)
        with tabs_rel[1]:
            df_d = df[df["Situação no plano"] == "Aguardando docs"]
            st.dataframe(df_d[["Nome", "E-mail corporativo", "Enviar no EB"]], use_container_width=True, hide_index=True)
        with tabs_rel[2]:
            df_dbl = df[df["Situação no plano"] == "Enviar à DBL"]
            st.dataframe(df_dbl[["Nome", "E-mail corporativo", "Enviar no EB"]], use_container_width=True, hide_index=True)
        with tabs_rel[3]:
            df_act = df[df["Situação no plano"] == "Aguardando DBL"]
            st.dataframe(df_act[["Nome", "E-mail corporativo", "Modelo de contrato"]], use_container_width=True, hide_index=True)

    # ----------------------------------------------------
    # 4. ABA AÇÕES
    # ----------------------------------------------------
    with aba_acoes:
        st.markdown("""
            <div style="background-color: #f1f3f5; padding: 12px; border-radius: 6px; border-left: 5px solid #404040; margin-bottom: 20px;">
                <span style="color: #404040; font-size: 14px;">Realize cadastros de benefícios, gere formulários e rascunhos de e-mail pré-preenchidos.</span>
            </div>
        """, unsafe_allow_html=True)

        # Divisão em 4 colunas igual ao DP
        c_cad, c_form, c_mail, c_div = st.columns(4)

        with c_cad:
            st.markdown("##### Cadastros")
            with st.expander("👤 Movimentações", expanded=False):
                st.caption("Em breve.")

        with c_form:
            st.markdown("##### Gerador de arquivos")
            
            # Expander 1: Inclusão PJ
            with st.expander("Inclusão PJ", expanded=False):
                if st.button("📄 Inclusão Subfatura", use_container_width=True):
                    modal_inclusao_subfatura(df)
                if st.button("📄 Termo Subestipulante", use_container_width=True):
                    modal_subestipulante(df)
            
            # Expander 2: Exclusão / Não Adesão
            with st.expander("Exclusão/Não Adesão PJ", expanded=False):
                if st.button("📄 Termo de Não Adesão", use_container_width=True):
                    modal_nao_adesao(df)
                if st.button("📄 Exclusão Subfatura", use_container_width=True):
                    modal_exclusao_subfatura()

        with c_mail:
            st.markdown("##### Mensagens")
            with st.expander("Comunicados", expanded=False):
                st.caption("Em breve.")

        with c_div:
            st.markdown("##### Diversos")
            with st.expander("Ferramentas", expanded=False):
                st.caption("Em breve.")
