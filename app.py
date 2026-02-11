import streamlit as st
import pandas as pd
import bcrypt
import altair as alt
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from docx import Document
from datetime import date

import bcrypt

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if "investidor_selecionado" not in st.session_state:
    st.session_state.investidor_selecionado = ""

def limpar_investidor():
    st.session_state.investidor_selecionado = ""

def formatar_cpf(valor):
    v = str(valor).replace(".0", "").zfill(11)
    if len(v) != 11:
        return ""
    return f"{v[:3]}.{v[3:6]}.{v[6:9]}-{v[9:]}"


def formatar_cnpj(valor):
    v = str(valor).replace(".0", "").zfill(14)
    if len(v) != 14:
        return ""
    return f"{v[:2]}.{v[2:5]}.{v[5:8]}/{v[8:12]}-{v[12:]}"

def render_table(df, *, dataframe=True, **kwargs):
    """
    Renderiza tabelas no Streamlit sem mostrar NaN / NaT / None,
    preservando os tipos originais do dataframe.
    """
    df_view = df.copy()

    # Substitui apenas para exibição
    df_view = df_view.where(pd.notna(df_view), "")

    if dataframe:
        st.dataframe(df_view, **kwargs)
    else:
        st.table(df_view)

def parse_data_br(coluna):
    return pd.to_datetime(coluna, dayfirst=True, errors="coerce")

from dateutil.relativedelta import relativedelta
import pandas as pd

def calcular_tempo_casa(data_inicio):
    if pd.isna(data_inicio):
        return ""

    hoje = pd.Timestamp.today().normalize()
    diff = relativedelta(hoje, data_inicio)

    return f"{diff.years} anos, {diff.months} meses e {diff.days} dias"

import unicodedata

def email_para_nome_arquivo(email):
    if not email:
        return ""

    email = unicodedata.normalize("NFKC", email)

    return (
        email
        .strip()
        .lower()
        .replace(" ", "")
    )

import re

def normalizar_cpf(cpf):
    if not cpf:
        return ""

    # remove tudo que não for número
    cpf = re.sub(r"\D", "", str(cpf))

    # garante 11 dígitos com zero à esquerda
    return cpf.zfill(11)

def gerar_hash_senha(senha):
    return bcrypt.hashpw(
        senha.encode("utf-8"),
        bcrypt.gensalt()
    ).decode("utf-8")

import pandas as pd

from docx import Document
from io import BytesIO

def gerar_docx_com_substituicoes(caminho_modelo, substituicoes):
    doc = Document(caminho_modelo)

    for paragrafo in doc.paragraphs:
        for run in paragrafo.runs:
            for chave, valor in substituicoes.items():
                if chave in run.text:
                    run.text = run.text.replace(chave, valor)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for run in paragrafo.runs:
                        for chave, valor in substituicoes.items():
                            if chave in run.text:
                                run.text = run.text.replace(chave, valor)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

from docx import Document

def substituir_runs_header_footer(doc, mapa):
    for section in doc.sections:
        # HEADER
        for p in section.header.paragraphs:
            for run in p.runs:
                for chave, valor in mapa.items():
                    if chave in run.text:
                        run.text = run.text.replace(chave, str(valor))

        # FOOTER
        for p in section.footer.paragraphs:
            for run in p.runs:
                for chave, valor in mapa.items():
                    if chave in run.text:
                        run.text = run.text.replace(chave, str(valor))

# --------------------------------------------------
# CONFIGURAÇÃO DA PÁGINA
# --------------------------------------------------

st.set_page_config(
    page_title="People | V4 Company",
    layout="wide",
    page_icon="LOGO VERMELHO.png"
)

# ==============================
# LOGIN
# ==============================
def verificar_senha(senha_digitada, senha_hash):
    return bcrypt.checkpw(
        senha_digitada.encode("utf-8"),
        senha_hash.encode("utf-8")
    )

import streamlit as st

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:

    st.title("🔐 Login")

    usuario = st.text_input("Usuário")
    senha = st.text_input("Senha", type="password")

    if st.button("Entrar"):

        users = st.secrets["users"]

        if usuario not in users:
            st.error("Usuário ou senha inválidos")
            st.stop()

        user = users[usuario]

        if not verificar_senha(senha, user["password"]):
            st.error("Usuário ou senha inválidos")
            st.stop()

        st.session_state.authenticated = True
        st.session_state.user_name = user["name"]
        st.rerun()

# ==============================
# LANDING PAGE (após login)
# ==============================

st.set_page_config(
    page_title="Dashboard People V4",
    layout="wide"
)

# Esconde menu padrão
hide_streamlit_style = """
    <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# CSS de centralização
st.markdown("""
    <style>
    .centered {
        display: flex;
        justify-content: center;
        align-items: center;
        height: 85vh;
        flex-direction: column;
        text-align: center;
    }
    .title {
        font-size: 60px;
        font-weight: bold;
    }
    .subtitle {
        font-size: 22px;
        color: gray;
        margin-top: -15px;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown(f"""
<div class="centered">
    <div class="title">People</div>
    <div class="subtitle">V4 Company</div>
</div>
""", unsafe_allow_html=True)
