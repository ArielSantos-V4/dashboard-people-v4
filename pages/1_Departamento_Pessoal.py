import streamlit as st
import pandas as pd
import bcrypt
import altair as alt
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from docx import Document
from datetime import date

import bcrypt

if "investidor_selecionado" not in st.session_state:
    st.session_state.investidor_selecionado = ""

def limpar_investidor():
    st.session_state.investidor_selecionado = ""

def formatar_cpf(valor):
    v = str(valor).replace(".0", "").zfill(11)
    if len(v) != 11:
        return ""
    return f"{v[:3]}.{v[3:6]}.{v[6:9]}-{v[9:]}"


def formatar_cnpj(valor):
    v = str(valor).replace(".0", "").zfill(14)
    if len(v) != 14:
        return ""
    return f"{v[:2]}.{v[2:5]}.{v[5:8]}/{v[8:12]}-{v[12:]}"

def render_table(df, *, dataframe=True, **kwargs):
    """
    Renderiza tabelas no Streamlit sem mostrar NaN / NaT / None,
    preservando os tipos originais do dataframe.
    """
    df_view = df.copy()

    # Substitui apenas para exibição
    df_view = df_view.where(pd.notna(df_view), "")

    if dataframe:
        st.dataframe(df_view, **kwargs)
    else:
        st.table(df_view)

def parse_data_br(coluna):
    return pd.to_datetime(coluna, dayfirst=True, errors="coerce")

from dateutil.relativedelta import relativedelta
import pandas as pd

def calcular_tempo_casa(data_inicio):
    if pd.isna(data_inicio):
        return ""

    hoje = pd.Timestamp.today().normalize()
    diff = relativedelta(hoje, data_inicio)

    return f"{diff.years} anos, {diff.months} meses e {diff.days} dias"

import unicodedata

def email_para_nome_arquivo(email):
    if not email:
        return ""

    email = unicodedata.normalize("NFKC", email)

    return (
        email
        .strip()
        .lower()
        .replace(" ", "")
    )

import re

def normalizar_cpf(cpf):
    if not cpf:
        return ""

    # remove tudo que não for número
    cpf = re.sub(r"\D", "", str(cpf))

    # garante 11 dígitos com zero à esquerda
    return cpf.zfill(11)

def gerar_hash_senha(senha):
    return bcrypt.hashpw(
        senha.encode("utf-8"),
        bcrypt.gensalt()
    ).decode("utf-8")

import pandas as pd

def gerar_alertas_investidor(linha):
    alertas = []

    # --- data de hoje (sem hora) ---
    hoje = pd.Timestamp.today().normalize()

    # --- status do plano ---
    status = str(linha["Situação no plano"]).strip()

    # =========================================================
    # ALERTA 1 — SOLICITAR DOCUMENTAÇÃO
    # status = Pendente → usa coluna "Solicitar documentação"
    # =========================================================
    data_solicitar = pd.to_datetime(
        linha["Solicitar documentação"],
        errors="coerce"
    )

    if status == "Pendente" and pd.notna(data_solicitar):
        dias = (data_solicitar - hoje).days

        if dias < 0:
            alertas.append((
                "error",
                "Plano de saúde e dental 🤕\n"
                "Solicitação de documentação em atraso. Verificar com urgência!"
            ))
        elif dias == 0:
            alertas.append((
                "warning",
                "Plano de saúde e dental ❤️‍🩹\n"
                "Hoje é a data limite para solicitar a documentação!"
            ))
        elif dias <= 15:
            alertas.append((
                "info",
                f"Plano de saúde e dental ❤️‍🩹\n"
                f"Faltam {dias} dias para solicitar a documentação ao investidor"
            ))

    # =========================================================
    # ALERTA 2 — ENVIAR NO EB
    # status = Aguardando docs → usa coluna "Enviar no EB"
    # =========================================================
    data_enviar_eb = pd.to_datetime(
        linha["Enviar no EB"],
        errors="coerce"
    )

    if status == "Aguardando docs" and pd.notna(data_enviar_eb):
        dias = (data_enviar_eb - hoje).days

        if dias < 0:
            alertas.append((
                "error",
                "Plano de saúde e dental 🤕\n"
                "Envio à EB em atraso. Verificar com urgência!"
            ))
        elif dias == 0:
            alertas.append((
                "warning",
                "Plano de saúde e dental ❤️‍🩹\n"
                "Hoje é a data limite para enviar à EB"
            ))
        elif dias <= 15:
            alertas.append((
                "info",
                f"Plano de saúde e dental ❤️‍🩹\n"
                f"Faltam {dias} dias para enviar à EB"
            ))

    if status == "Aguardando DBL":
        alertas.append(("info",
            "Plano de saúde e dental quase prontos! 🤩"
            "Acompanhar movimentação no portal EB"
        ))
    
    # -------------------------
    # ALERTA — Aniversário
    # -------------------------
    nascimento_raw = linha.get("Data de nascimento", "")
    
    nascimento = pd.to_datetime(
        nascimento_raw,
        errors="coerce",
        dayfirst=True
    )
    
    if pd.notna(nascimento):
        nascimento = pd.Timestamp(nascimento).normalize()
    
        if nascimento.month == hoje.month:
            if nascimento.day == hoje.day:
                alertas.append((
                    "info",
                    "Lembrete de Aniversário! 🎉\n"
                    "HOJE é aniversário do investidor!!"
                ))
            else:
                alertas.append((
                    "info",
                    "Lembrete de Aniversário! 🎉\n"
                    "Este investidor faz aniversário neste mês"
                ))

    # -------------------------
    # ALERTA 3 — Contrato
    # -------------------------
    fim_contrato_raw = linha.get("Térm previsto", "")

    fim_contrato = pd.to_datetime(
        fim_contrato_raw,
        errors="coerce",
        dayfirst=True
    )
    
    if pd.notna(fim_contrato):
        fim_contrato = pd.Timestamp(fim_contrato).normalize()
        dias = (fim_contrato - hoje).days
    
    if pd.notna(fim_contrato):
        dias = (fim_contrato - hoje).days

        if dias < 0:
            alertas.append(("error",
                "Contrato vencido! 🚨"
                "Verificar com urgência!"
            ))
        elif dias <= 30:
            alertas.append(("warning",
                f"Alerta! ⚠️"
                f"O contrato se encerra em {dias} dia(s)."
            ))

    # -------------------------
    # ALERTA 4 — MEI
    # -------------------------
    if linha.get("Modalidade PJ", "") == "MEI":
        alertas.append(("warning",
            "Atenção! Investidor ainda se encontra na modalidade MEI 😬"
        ))

    return alertas

st.markdown("""
<style>
/* Modal específico da consulta individual */
div[role="dialog"]:has(.modal-investidor) {
    width: 95vw !important;
    max-width: 95vw !important;
}

/* Altura maior (opcional) */
div[role="dialog"]:has(.modal-investidor) > div {
    max-height: 90vh !important;
}
    
</style>
""", unsafe_allow_html=True)

from docx import Document
from io import BytesIO

def gerar_docx_com_substituicoes(caminho_modelo, substituicoes):
    doc = Document(caminho_modelo)

    for paragrafo in doc.paragraphs:
        for run in paragrafo.runs:
            for chave, valor in substituicoes.items():
                if chave in run.text:
                    run.text = run.text.replace(chave, valor)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for run in paragrafo.runs:
                        for chave, valor in substituicoes.items():
                            if chave in run.text:
                                run.text = run.text.replace(chave, valor)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

from docx import Document

def gerar_vale_transporte(dados):
    # Abre o documento ORIGINAL (modelo)
    doc = Document("modelo_vale_transporte.docx")

    for p in doc.paragraphs:
        if "{{NOME}}" in p.text:
            p.text = p.text.replace("{{NOME}}", dados["nome"])

        if "{{CPF}}" in p.text:
            p.text = p.text.replace("{{CPF}}", dados["cpf"])

        if "{{VALOR}}" in p.text:
            p.text = p.text.replace("{{VALOR}}", dados["valor"])

    doc.save("vale_transporte_final.docx")

def substituir_runs_header_footer(doc, mapa):
    for section in doc.sections:
        # HEADER
        for p in section.header.paragraphs:
            for run in p.runs:
                for chave, valor in mapa.items():
                    if chave in run.text:
                        run.text = run.text.replace(chave, str(valor))

        # FOOTER
        for p in section.footer.paragraphs:
            for run in p.runs:
                for chave, valor in mapa.items():
                    if chave in run.text:
                        run.text = run.text.replace(chave, str(valor))

# --------------------------------------------------
# CONFIGURAÇÃO DA PÁGINA
# --------------------------------------------------

st.set_page_config(
    page_title="People | V4 Company",
    layout="wide",
    page_icon="LOGO VERMELHO.png"
)


# --------------------------------------------------
# ABAS
# --------------------------------------------------
aba_dashboard, aba_relatorios = st.tabs([
    "📊 Dashboard",
    "📄 Relatórios",
])

# --------------------------------------------------
# ABA DASHBOARD
# --------------------------------------------------

with aba_dashboard:
    
    # --------------------------------------------------
    # ESTILO
    # --------------------------------------------------
    st.markdown("""
    <style>
    
    /* =========================
       CONSULTA INDIVIDUAL — COMPACTAÇÃO REAL
       ========================= */
    
    /* Títulos das seções */
    h5 {
        font-size: 20px !important;
        margin-top: 6px !important;
        margin-bottom: 2px !important;
    }
    
    /* Label */
    label {
        font-size: 10px !important;
        margin-bottom: 0px !important;
        color: #bdbdbd !important;
    }
    
    /* 🔥 CONTAINER DO INPUT (o retângulo) */
    div[data-testid="stTextInput"] {
        height: 30px !important;
    }
    
    /* 🔥 INPUT REAL */
    div[data-testid="stTextInput"] input {
        height: 40px !important;
        padding: 10px 10px !important;
        font-size: 12px !important;
        line-height: 0px !important; /* 👈 CENTRALIZA O TEXTO */
    }
    
    /* Remove espaço entre campos */
    div[data-testid="stTextInput"] {
        margin-bottom: 25px !important;
    }
    
    /* Remove respiro extra das colunas */
    div[data-testid="column"] {
        padding-top: 5px !important;
        padding-bottom: 0px !important;
    }
    
    /* Benefícios */
    .espaco-beneficio {
        margin-top: 15px;
        margin-bottom: 4px;
    }
    
    </style>
    """, unsafe_allow_html=True)
           
    # --------------------------------------------------
    # GOOGLE SHEETS
    # --------------------------------------------------
    import gspread
    from google.oauth2.service_account import Credentials
    
    @st.cache_data(ttl=600)
    def load_google_sheet():
        creds = Credentials.from_service_account_info(
            st.secrets["gcp_service_account"],
            scopes=["https://www.googleapis.com/auth/spreadsheets.readonly"]
        )
    
        client = gspread.authorize(creds)
    
        sheet = client.open_by_key("13EPwhiXgh8BkbhyrEy2aCy3cv1O8npxJ_hA-HmLZ-pY")
        worksheet = sheet.get_worksheet(5)
    
        data = worksheet.get_all_records()
        return pd.DataFrame(data)

    
    # --------------------------------------------------
    # LOAD + ORGANIZAÇÃO
    # --------------------------------------------------
    df = load_google_sheet()

    # Padronização de colunas
    df = df.rename(columns={
        "Nome completo": "Nome",
        "Data Início": "Início na V4",
        "Término contrato previsto": "Térm previsto",
        "Ativo no plano": "Situação no plano"
    })

    df.columns = (
        df.columns
        .astype(str)
        .str.replace("\u00a0", " ", regex=False)
        .str.strip()
    )

    if "Nome" not in df.columns:
        st.error("❌ A coluna 'Nome' não foi encontrada na planilha.")
        st.write("Colunas disponíveis:", df.columns.tolist())
        st.stop()

    df = df.sort_values(df.columns[0]).reset_index(drop=True)
    
    # 👇 AQUI É O LUGAR CERTO
    df = df.fillna("")
            
    # --------------------------------------------------
    # CONVERSÃO CORRETA (DAYFIRST)
    # --------------------------------------------------
    
    # BACKUP TEXTO ORIGINAL
    df["Início na V4_raw"] = df["Início na V4"]
    df["Data de nascimento_raw"] = df["Data de nascimento"]
    df["Data do contrato_raw"] = df.iloc[:, 12]
    df["Térm previsto_raw"] = df.iloc[:, 6]
    
    # DATETIME (PARA CÁLCULOS)
    df["Início na V4_dt"] = parse_data_br(df["Início na V4_raw"])
    df["Data de nascimento_dt"] = parse_data_br(df["Data de nascimento_raw"])
    df["Data do contrato_dt"] = parse_data_br(df["Data do contrato_raw"])
    df["Térm previsto_dt"] = parse_data_br(df["Térm previsto_raw"])
    
    # TEXTO FINAL (EXIBIÇÃO)
    df["Início na V4"] = df["Início na V4_dt"].dt.strftime("%d/%m/%Y").fillna("")
    df["Data de nascimento"] = df["Data de nascimento_dt"].dt.strftime("%d/%m/%Y").fillna("")
    df["Data do contrato"] = df["Data do contrato_dt"].dt.strftime("%d/%m/%Y").fillna("")
    
    # Térm previsto: data vira data, texto continua texto
    df["Térm previsto"] = df["Térm previsto_raw"].where(
        df["Térm previsto_dt"].isna(),
        df["Térm previsto_dt"].dt.strftime("%d/%m/%Y")
    )

    # --------------------------------------------------
    # SIDEBAR
    # --------------------------------------------------
    st.sidebar.success(
        f"Bem-vindo(a), {st.session_state.get('user_name', 'Usuário')}"
    )
    
    if st.sidebar.button("🔄 Atualizar dados"):
        st.cache_data.clear()
        st.rerun()
    
    if st.sidebar.button("Logout"):
        st.session_state.authenticated = False
        st.rerun()

    st.sidebar.divider()

    # --------------------------------------------------
    # TOPO
    # --------------------------------------------------
    col_logo, col_title = st.columns([1, 6])
    
    with col_logo:
        st.image("LOGO VERMELHO.png", width=120)
        
    with col_title:
        st.markdown("<h1>Dashboard People</h1><h3 style='color:#ccc;'>V4 Company</h3>", unsafe_allow_html=True)
    
    st.markdown("---")

    # --------------------------------------------------
    # CONSULTA INDIVIDUAL
    # --------------------------------------------------
    st.markdown("""
    <style>
        /* Remove header APENAS do modal da consulta individual */
        div[role="dialog"]:has(.modal-investidor) > div > header {
            display: none !important;
        }
        
        /* Remove o espaço do header só nesse modal */
        div[role="dialog"]:has(.modal-investidor) > div {
            padding-top: 0px !important;
        }
    
    </style>
    """, unsafe_allow_html=True)


    @st.dialog(" ")
    def modal_consulta_investidor(df_consulta, nome):
        st.markdown('<div class="modal-investidor">', unsafe_allow_html=True)

        linha = df_consulta[df_consulta["Nome"] == nome].iloc[0]
             
        col1, col2, col3 = st.columns([3, 3, 2])
            
        # -------------------------
        # COLUNA 1 — PROFISSIONAL
        # -------------------------
        with col1:
            st.markdown("##### 📌 Dados profissionais")
            
            bp = str(linha["BP"]).replace(".0", "")
            matricula = str(linha["Matrícula"]).replace(".0", "").zfill(6)
         
            a1, a2, a3 = st.columns(3)
            a1.text_input("BP", bp, disabled=True)
            a2.text_input("Matrícula", matricula, disabled=True)
            a3.text_input("Situação", linha["Situação"], disabled=True)
        
            a4, a5, a6 = st.columns(3)
            a4.text_input("Data do contrato", linha["Data do contrato"], disabled=True)
            a5.text_input("Término previsto", linha["Térm previsto"], disabled=True)
            a6.text_input("Modelo contrato", linha["Modelo de contrato"], disabled=True)
            
            tempo_casa = ""
            if linha["Início na V4"] != "":
                delta = datetime.today() - linha["Início na V4_dt"]
                anos = delta.days // 365
                meses = (delta.days % 365) // 30
                dias = (delta.days % 365) % 30
                tempo_casa = f"{anos} anos, {meses} meses e {dias} dias"
            
            a7, a8 = st.columns([1, 2])
            a7.text_input("Início na V4", linha["Início na V4"], disabled=True)
            a8.text_input("Tempo de casa", tempo_casa, disabled=True)
       
            a9, a10 = st.columns([3, 1])
            a9.text_input("Unidade / Atuação", linha["Unidade/Atuação"], disabled=True)
            a10.text_input("Modalidade PJ", linha["Modalidade PJ"], disabled=True)
           
            st.text_input("E-mail corporativo", linha["E-mail corporativo"], disabled=True)
          
            cnpj = formatar_cnpj(linha["CNPJ"])
        
            a11, a12 = st.columns(2)
            a11.text_input("CNPJ", cnpj, disabled=True)
            a12.text_input("Razão social", linha["Razão social"], disabled=True)
        
        
        
            a13, a14 = st.columns([3, 1])
            a13.text_input("Cargo", linha["Cargo"], disabled=True)
            a14.text_input("Remuneração", linha["Remuneração"], disabled=True)
            
            a15, a16 = st.columns([1, 3])
            a15.text_input("CBO", linha["CBO"], disabled=True)
            a16.text_input("Descrição CBO", linha["Descrição CBO"], disabled=True)
        
        
        # -------------------------
        # COLUNA 2 — ADMIN / PESSOAL
        # -------------------------
        with col2:
            st.markdown("##### 🧾 Centro de custo")
    
            # Centro de custo (código menor / descrição maior)
            codigo_cc = str(linha["Código CC"]).replace(".0", "")
    
            b1, b2 = st.columns([1, 3])
            b1.text_input("Código CC", codigo_cc, disabled=True)
            b2.text_input("Descrição CC", linha["Descrição CC"], disabled=True)
    
    
            b3, b4 = st.columns(2)
            b3.text_input("Senioridade", linha["Senioridade"], disabled=True)
            b4.text_input("Conta contábil", linha["Conta contábil"], disabled=True)
    
            st.text_input("Liderança direta", linha["Liderança direta"], disabled=True)
    
            st.markdown("##### 👤 Dados pessoais")
    
            cpf = str(linha["CPF"]).replace(".0", "")
    
            b5, b6, b7 = st.columns(3)
            cpf = formatar_cpf(linha["CPF"])
            b5.text_input("CPF", cpf, disabled=True)
            b6.text_input("Nascimento", linha["Data de nascimento"], disabled=True)
  
            idade = ""
            if linha["Data de nascimento"] != "":
                idade = int((datetime.today() - pd.to_datetime(linha["Data de nascimento"])).days / 365.25)
                idade = f"{idade} anos"
            b7.text_input("Idade", idade, disabled=True)
    
            b8, b9 = st.columns(2)
            b8.text_input("CEP", linha["CEP"], disabled=True)
            b9.text_input("Escolaridade", linha["Escolaridade"], disabled=True)
    
            st.text_input("Telefone pessoal", linha["Telefone pessoal"], disabled=True)
            st.text_input("E-mail pessoal", linha["E-mail pessoal"], disabled=True)
    
        # -------------------------
        # COLUNA 3 — FOTO / BENEFÍCIOS / LINK
        # -------------------------
        with col3:
            st.markdown("##### 🖼️ Foto")
            if linha["Foto"]:
                st.markdown(
                    f"""
                    <div style="display:flex; justify-content:center;">
                        <img src="{linha['Foto']}" width="160">
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            else:
                st.info("Sem foto")
    
            st.markdown("##### 🎁 Benefícios")
    
            st.text_input("Situação no plano", linha["Situação no plano"], disabled=True)
    
            carteira_med = str(linha["Carteirinha médico"]).replace(".0", "")
            carteira_odo = str(linha["Carteirinha odonto"]).replace(".0", "")

            m1, m2 = st.columns(2)
            m1.text_input("Plano médico", linha["Operadora Médico"], disabled=True)
            m2.text_input("Carteirinha médico", carteira_med, disabled=True)
    
            st.markdown('<div class="espaco-beneficio"></div>', unsafe_allow_html=True)
    
            o1, o2 = st.columns(2)
            o1.text_input("Plano odonto", linha["Operadora Odonto"], disabled=True)
            o2.text_input("Carteirinha odonto", carteira_odo, disabled=True)
    
            col_link, col_alertas = st.columns([1, 3])
            
            # --- LINK DRIVE ---
            with col_link:
                st.markdown("##### 🔗 Link")
                if linha["Link Drive"]:
                    st.link_button("Drive", linha["Link Drive"])
                else:
                    st.caption("Sem link de Drive")
            
            # --- ALERTAS ---
            with col_alertas:
                st.markdown("##### ⚠️ Alertas")
                alertas = st.session_state.get("alertas_atuais", [])
            
                if alertas:
                    with st.container(height=100, border=True):
                        for tipo, mensagem in alertas:
                            if tipo == "error":
                                st.error(mensagem)
                            elif tipo == "warning":
                                st.warning(mensagem)
                            else:
                                st.info(mensagem)

        
        st.markdown('</div>', unsafe_allow_html=True)        
    
    st.subheader("🔎 Consulta individual do investidor")
        
    df_consulta = df.fillna("")
    lista_nomes = sorted(df_consulta["Nome"].unique())
        
    with st.form("form_consulta_investidor", clear_on_submit=False):
        c1, c2, c3 = st.columns([6, 1, 1])
    
        with c1:
            nome = st.selectbox(
                "Selecione o investidor",
                ["Selecione um investidor..."] + lista_nomes,
                key="investidor_selecionado",
                label_visibility="collapsed"
            )
    
        with c2:
            consultar = st.form_submit_button("🔍 Consultar")
    
        with c3:
            limpar = st.form_submit_button("Limpar")
    
        if consultar and st.session_state.investidor_selecionado != "Selecione um investidor...":

            # pega a linha do investidor selecionado
            linha = df_consulta[
                df_consulta["Nome"] == st.session_state.investidor_selecionado
            ].iloc[0]
        
            # gera e salva os alertas
            st.session_state.alertas_atuais = gerar_alertas_investidor(linha)
        
            # abre o modal
            modal_consulta_investidor(
                df_consulta,
                st.session_state.investidor_selecionado
            )
           
        if limpar:
            limpar_investidor()
            st.session_state.abrir_modal_investidor = False
                                   
    # --------------------------------------------------
    # FORMAT TABELA
    # --------------------------------------------------
    
    def limpar_numero(valor):
        if valor == "" or pd.isna(valor):
            return ""
        return str(valor).replace(".0", "").strip()
    
    
    def formatar_cpf(valor):
        v = limpar_numero(valor)
        if len(v) == 11:
            return f"{v[0:3]}.{v[3:6]}.{v[6:9]}-{v[9:11]}"
        return v
    
    
    def formatar_cnpj(valor):
        v = limpar_numero(valor)
        if len(v) == 14:
            return f"{v[0:2]}.{v[2:5]}.{v[5:8]}/{v[8:12]}-{v[12:14]}"
        return v
    
    
    def formatar_matricula(valor):
        v = limpar_numero(valor)
        if v.isdigit():
            return v.zfill(6)
        return v
    
    # --------------------------------------------------
    # TABELA
    # --------------------------------------------------
    st.markdown("---")
    st.markdown("### 📋 Base de investidores")
    
    busca = st.text_input(
        "Buscar na tabela",
        placeholder="🔍 Buscar na tabela...",
        label_visibility="collapsed"
    )
    
    
    df_tabela = df.copy()

    df_tabela["Data de nascimento"] = df_tabela["Data de nascimento"]
    df_tabela["Data do contrato"] = df_tabela["Data do contrato"]
    df_tabela["Início na V4"] = df_tabela["Início na V4"]

    # Datas exibidas
    df_tabela["Término do contrato"] = df_tabela["Térm previsto"]
    df_tabela["Data de início"] = df_tabela["Início na V4"]
    
    # Limpeza de campos com .0
    df_tabela["BP"] = df_tabela["BP"].apply(limpar_numero)
    df_tabela["Código CC"] = df_tabela["Código CC"].apply(limpar_numero)
    for col in ["Carteirinha médico", "Carteirinha odonto"]:
        if col in df_tabela.columns:
            df_tabela[col] = df_tabela[col].apply(limpar_numero)
    
    # Matrícula com 6 dígitos
    df_tabela["Matrícula"] = df_tabela["Matrícula"].apply(formatar_matricula)
    
    # CPF e CNPJ formatados
    df_tabela["CPF"] = df_tabela["CPF"].apply(formatar_cpf)
    df_tabela["CNPJ"] = df_tabela["CNPJ"].apply(formatar_cnpj)
    
    
    if busca:
        df_tabela = df_tabela[
            df_tabela.astype(str)
            .apply(lambda x: x.str.contains(busca, case=False).any(), axis=1)
        ]
        
    df_tabela.insert(
        df_tabela.columns.get_loc("Nome") + 1,
        "Início na V4",
        df_tabela.pop("Início na V4")
    )
    

    st.dataframe(
        df_tabela.drop(
            columns=[c for c in df_tabela.columns if c.endswith("_raw") or c.endswith("_dt")],
            errors="ignore"
        ),
        use_container_width=True,
        hide_index=True
    )

     
    # --------------------------------------------------
    # KPIs
    # --------------------------------------------------
    st.markdown("---")
    hoje = datetime.today()
    prox_30_dias = hoje + timedelta(days=30)
    
    headcount = len(df)
    contratos_vencer = df[
        df["Térm previsto_dt"].notna() &
        (df["Térm previsto_dt"] <= prox_30_dias)
    ]
    
    contratos_vencidos = df[
        df["Térm previsto_dt"].notna() &
        (df["Térm previsto_dt"] < hoje)
    ]
    
    pj = len(df[df["Modelo de contrato"] == "PJ"])
    clt = len(df[df["Modelo de contrato"] == "CLT"])
    estagio = len(df[df["Modelo de contrato"] == "Estágio"])
    
    df_adm = df[df["Início na V4_dt"].notna()]

    media_admissoes = (
        df_adm
        .groupby(df_adm["Início na V4_dt"].dt.to_period("M"))
        .size()
        .mean()
    )
    
    # --------------------------------------------------
    # KPIs VISUAIS
    # --------------------------------------------------
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Headcount", headcount)
    c2.metric("Contratos vencendo (30 dias)", len(contratos_vencer))
    c3.metric("Contratos vencidos", len(contratos_vencidos))
    c4.metric("PJ / CLT / Estágio", f"{pj} / {clt} / {estagio}")
    c5.metric("Média admissões / mês", f"{media_admissoes:.1f}")
    
    st.markdown("---")
    
    # --------------------------------------------------
    # GRÁFICOS
    # --------------------------------------------------
    g1, g2 = st.columns(2)
    
    with g1:
        st.subheader("📃 Modelo de contrato")
        contrato_df = df["Modelo de contrato"].value_counts().reset_index()
        contrato_df.columns = ["Modelo", "Quantidade"]
    
        st.altair_chart(
            alt.Chart(contrato_df)
            .mark_arc(innerRadius=60)
            .encode(
                theta="Quantidade:Q",
                color=alt.Color("Modelo:N", scale=alt.Scale(range=["#E30613", "#B0000A", "#FF4C4C"])),
                tooltip=["Modelo", "Quantidade"]
            ),
            use_container_width=True
        )
    
    with g2:
        st.subheader("📍 Local de atuação")
        local_df = df["Unidade/Atuação"].value_counts().reset_index()
        local_df.columns = ["Local", "Quantidade"]
    
        st.altair_chart(
            alt.Chart(local_df)
            .mark_bar(color="#E30613")
            .encode(
                x=alt.X("Local:N", sort="-y", axis=alt.Axis(labelAngle=-30)),
                y="Quantidade:Q",
                tooltip=["Local", "Quantidade"]
            ),
            use_container_width=True
        )
    
    # --------------------------------------------------
    # ADMISSÕES
    # --------------------------------------------------
    st.subheader("📈 Admissões por mês")
    
    adm_mes = (
        df_adm.assign(Mes=df_adm["Início na V4_dt"].dt.strftime("%b/%Y"))
        .groupby("Mes")
        .size()
        .reset_index(name="Quantidade")
    )
    
    st.altair_chart(
        alt.Chart(adm_mes)
        .mark_line(color="#E30613", point=True)
        .encode(x="Mes:N", y="Quantidade:Q", tooltip=["Mes", "Quantidade"]),
        use_container_width=True
    )
    

# --------------------------------------------------
# ABA RELATÓRIOS
# --------------------------------------------------
with aba_relatorios:

    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)

    # --------------------------------------------------
    # TOPO
    # --------------------------------------------------
    col_logo, col_title = st.columns([1, 6])

    with col_logo:
        st.image("LOGO VERMELHO.png", width=120)

    with col_title:
        st.markdown(
            "<h1>Análises & Relatórios</h1><h3 style='color:#ccc;'>V4 Company</h3>",
            unsafe_allow_html=True
        )

    st.markdown("---")

    # --------------------------------------------------
    # LAYOUT PRINCIPAL — RELATÓRIOS
    # --------------------------------------------------
    col_relatorios, col_divisor, col_acoes = st.columns([7, 0.1, 3])
    
    with col_divisor:
        st.markdown(
            """
            <div style="
                height: 100%;
                border-left: 1px solid #e0e0e0;
                margin: 0 auto;
            "></div>
            """,
            unsafe_allow_html=True
        )


    # --------------------------------------------------
    # COLUNA ESQUERDA — RELATÓRIOS
    # --------------------------------------------------
    with col_relatorios:

        st.markdown("## 📊 Relatórios Principais")

        # -------------------------------
        # ANIVERSARIANTES DO MÊS
        # -------------------------------
        
        with st.expander("🎉 Aniversariantes do mês", expanded=False):
        
            meses = {
                1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
                5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
                9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
            }
        
            mes_atual = datetime.today().month
        
            mes_selecionado = st.selectbox(
                "Mês",
                options=list(meses.keys()),
                format_func=lambda x: meses[x],
                index=mes_atual - 1
            )
        
            df_aniversario = df.copy()
        
            df_aniversario = df[
                df["Data de nascimento_dt"].dt.month == mes_selecionado
            ]

            df_check = df.copy()

            df_check["Data de nascimento_raw"] = df_check["Data de nascimento"]
            
            df_check["Data de nascimento_dt"] = pd.to_datetime(
                df_check["Data de nascimento_raw"],
                dayfirst=True,
                errors="coerce"
            )
            
            df_invalidos = df_check[df_check["Data de nascimento_dt"].isna()]


            # 🔔 LISTAR PESSOAS COM DATA INVÁLIDA
            if not df_invalidos.empty:
                col_warn, col_link = st.columns([5, 2])
            
                with col_warn:
                    st.warning(f"⚠️ {len(df_invalidos)} pessoas com data de nascimento inválida")
            
                with col_link:
                    with st.popover("👀 Ver aqui"):
                        df_invalidos_view = df_invalidos[
                            ["Nome", "Data de nascimento_raw"]
                        ].reset_index(drop=True)
            
                        st.table(df_invalidos_view)
        
            if df_aniversario.empty:
                st.info("Nenhum aniversariante neste mês 🎈")
            else:
                ano_atual = datetime.today().year
        
                df_aniversario["Nascimento"] = df_aniversario["Data de nascimento_dt"].dt.strftime("%d/%m/%Y")
        
                df_aniversario["Idade que completa"] = (
                    ano_atual - df_aniversario["Data de nascimento_dt"].dt.year
                ).astype(int).astype(str) + " anos"
        
                df_aniversario["Dia"] = df_aniversario["Data de nascimento_dt"].dt.day
        
                df_final = df_aniversario[
                    ["Nome", "E-mail corporativo", "Nascimento", "Idade que completa", "Dia"]
                ].sort_values("Dia")
        
                # 🔥 remove índice visual
                df_final = df_final.reset_index(drop=True)
                df_final.index = [""] * len(df_final)
        
                render_table(
                    df_final.drop(columns=["Dia"]),
                    use_container_width=True,
                    hide_index=True,
                    column_config={
                        "Nascimento": st.column_config.TextColumn(
                            "Nascimento",
                            width="small"
                        ),
                        "Idade que completa": st.column_config.TextColumn(
                            "Idade que completa",
                            width="small"
                        ),
                        "Nome": st.column_config.TextColumn(
                            "Nome",
                            width="large"
                        ),
                        "E-mail corporativo": st.column_config.TextColumn(
                            "E-mail corporativo",
                            width="large"
                        ),
                    }
                )


        # -------------------------------
        # VENCIMENTO / TÉRMINO PREVISTO
        # -------------------------------
        
        with st.expander("⏰ Contratos a vencer", expanded=False):
        
            col1, col2 = st.columns(2)
        
            with col1:
                data_inicio = st.date_input(
                    "Data inicial",
                    value=datetime.today().date(),
                    format="DD/MM/YYYY"
                )
            
            with col2:
                data_fim = st.date_input(
                    "Data final",
                    value=datetime.today().date() + relativedelta(months=3),
                    format="DD/MM/YYYY"
                )

        
            # Garante datetime
            df["Térm previsto_dt"] = pd.to_datetime(
                df["Térm previsto"],
                dayfirst=True,
                errors="coerce"
            )
            
            # Converte datas do Streamlit para Timestamp
            data_inicio_ts = pd.Timestamp(data_inicio)
            data_fim_ts = pd.Timestamp(data_fim)
            
            # Filtro correto
            df_vencimento = df[
                df["Térm previsto_dt"].notna() &
                (df["Térm previsto_dt"] >= data_inicio_ts) &
                (df["Térm previsto_dt"] <= data_fim_ts)
            ]
        
            # 🔹 ordena ANTES de cortar colunas
            df_vencimento = df_vencimento.sort_values(
                "Térm previsto_dt",
                na_position="last"
            )
        
            if df_vencimento.empty:
                st.info("Nenhum contrato vencendo no período selecionado ⏳")
            else:
                # 🔹 formata data apenas para exibição
                df_vencimento["Térm previsto"] = (
                    df_vencimento["Térm previsto_dt"]
                    .dt.strftime("%d/%m/%Y")
                    .fillna("")
                )
        
                df_final = df_vencimento[
                    [
                        "Nome",
                        "E-mail corporativo",
                        "Térm previsto"
                    ]
                ].reset_index(drop=True)
        
                df_final.index = [""] * len(df_final)
        
                render_table(
                    df_final,
                    use_container_width=True,
                    hide_index=True,
                    column_config={
                        "Nome": st.column_config.TextColumn(
                            "Nome",
                            width="large"
                        ),
                        "E-mail corporativo": st.column_config.TextColumn(
                            "E-mail corporativo",
                            width="large"
                        ),
                        "Término previsto": st.column_config.TextColumn(
                            "Térm previsto",
                            width="small"
                        ),
                    }
                )

        # -------------------------------
        # INVESTIDORES MEI
        # -------------------------------
        with st.expander("💼 Investidores MEI", expanded=False):

            # Verifica se a coluna Modalidade PJ existe
            if "Modalidade PJ" not in df.columns:
                st.warning("Coluna 'Modalidade PJ' não encontrada no DataFrame.")
        
            else:
                # Filtra apenas MEI
                df_mei = df[
                    df["Modalidade PJ"]
                    .astype(str)
                    .str.upper()
                    .str.contains("MEI", na=False)
                ]
        
                if df_mei.empty:
                    st.info("Nenhum investidor MEI encontrado.")
        
                else:

                    # 🔔 ALERTA – TOTAL DE INVESTIDORES MEI
                    total_mei = len(df_mei)
                    
                    st.warning(
                        f"⚠️ Temos **{total_mei} investidores na modalidade MEI** que precisam regularizar a forma de contratação."
                    )

                    # 🔹 MAPEAMENTO SEGURO DE COLUNAS
                    colunas_map = {
                        "Nome": None,
                        "Email Corporativo": None,
                        "Data do contrato": None,
                        "Modalidade PJ": None,
                    }
                    
                    for col in df_mei.columns:
                        c = col.strip().lower()
                    
                        if c == "nome":
                            colunas_map["Nome"] = col
                    
                        elif "mail" in c:
                            colunas_map["Email Corporativo"] = col
                    
                        elif any(x in c for x in ["contrato", "admiss"]):
                            colunas_map["Data do contrato"] = col
                    
                        elif "modalidade" in c:
                            colunas_map["Modalidade PJ"] = col

        
                    # Remove colunas não encontradas
                    colunas_validas = {
                        k: v for k, v in colunas_map.items() if v is not None
                    }
        
                    df_mei_final = df_mei[list(colunas_validas.values())].copy()
                    df_mei_final.columns = list(colunas_validas.keys())
        
                    # Formata data do contrato
                    if "Data do contrato" in df_mei_final.columns:
                        df_mei_final["Data do contrato"] = pd.to_datetime(
                            df_mei_final["Data do contrato"],
                            errors="coerce"
                        ).dt.strftime("%d/%m/%Y")
        
                    st.dataframe(
                        df_mei_final,
                        use_container_width=True,
                        hide_index=True
                    )

        # ==============================
        # RELATÓRIO — TEMPO DE CASA
        # ==============================
        with st.expander("⏳ Tempo de Casa", expanded=False):
            
            from dateutil.relativedelta import relativedelta
            
            def calcular_tempo_casa(data_inicio):
                if pd.isna(data_inicio):
                    return ""
                hoje = pd.Timestamp.today().normalize()
                diff = relativedelta(hoje, data_inicio)
                return f"{diff.years} anos, {diff.months} meses e {diff.days} dias"
            
            
            df_relatorio_tempo = df.copy()
            
            # 🔹 MAPEAMENTO SEGURO DA DATA DE INÍCIO
            col_inicio = None
            for col in df_relatorio_tempo.columns:
                c = col.lower().strip()
                if "início" in c or "inicio" in c or "admiss" in c or "contrato" in c:
                    col_inicio = col
                    break
            
            if col_inicio is None:
                st.error("Coluna de início não encontrada.")
            else:
                df_relatorio_tempo["Inicio_dt"] = pd.to_datetime(
                    df_relatorio_tempo[col_inicio],
                    dayfirst=True,
                    errors="coerce"
                )
            
                df_relatorio_tempo["Tempo de casa"] = df_relatorio_tempo["Inicio_dt"].apply(
                    calcular_tempo_casa
                )
            
                # 🔎 FILTRO
                min_anos = st.selectbox(
                    "Tempo mínimo de casa (anos)",
                    [0, 1, 2, 3, 4, 5],
                    index=0
                )
            
                if min_anos > 0:
                    hoje = pd.Timestamp.today().normalize()
                    df_relatorio_tempo = df_relatorio_tempo[
                        (hoje - df_relatorio_tempo["Inicio_dt"]).dt.days >= min_anos * 365
                    ]
            
                df_final = df_relatorio_tempo[
                    [
                        "Nome",
                        "E-mail corporativo",
                        col_inicio,
                        "Remuneração",
                        "Tempo de casa"
                    ]
                ].rename(columns={col_inicio: "Início na V4"})
            
                st.dataframe(df_final, use_container_width=True, hide_index=True)


    # --------------------------------------------------
    # COLUNA DIREITA — AÇÕES
    # --------------------------------------------------
    with col_acoes:

        st.markdown("## ⚙️ Ações")

        # ---------------------------------
        # BOTÃO – TÍTULO DE DOC PARA AUTOMAÇÃO
        # ---------------------------------
        
        def limpar_titulo():
            st.session_state["titulo_doc"] = ""
            st.session_state.pop("titulo_gerado", None)
        
        
        @st.dialog("📝 Gerador de título para automação")
        def modal_titulo_doc():
        
            # ---------- CAMPO TÍTULO + BOTÃO LIMPAR ----------
            col_input, col_clear = st.columns([5, 1])
        
            with col_input:
                st.text_input(
                    "Título original do arquivo",
                    placeholder="Cole aqui o título do arquivo",
                    key="titulo_doc"
                )
        
            with col_clear:
                st.markdown("<div style='height:23px'></div>", unsafe_allow_html=True)
                st.button(
                    "❌",
                    help="Limpar título",
                    on_click=limpar_titulo
                )
        
            # ---------- SELECT DE NOMES (ALFABÉTICO / EM BRANCO) ----------
            lista_nomes = sorted(df["Nome"].dropna().unique())
        
            st.selectbox(
                "Selecione o investidor",
                options=[""] + lista_nomes,
                index=0,
                key="nome_selecionado",
                placeholder="Digite ou selecione um nome"
            )

            if st.button("✅ Gerar", use_container_width=True):
                gerar = True
       
                titulo_doc = st.session_state.get("titulo_doc", "")
                nome_selecionado = st.session_state.get("nome_selecionado", "")
        
                if not nome_selecionado or not titulo_doc:
                    st.warning("Selecione um nome e informe o título do arquivo.")
                    return
        
                dados_filtrados = df[df["Nome"] == nome_selecionado]
       
                if dados_filtrados.empty:
                    st.error("Não foi possível localizar os dados dessa pessoa.")
                    return
        
                dados = dados_filtrados.iloc[0]
     
                cpf_limpo = (
                    str(dados.get("CPF", ""))
                    .replace(".", "")
                    .replace("-", "")
                    .replace("/", "")
                    .zfill(11)   # 👈 garante 11 dígitos, incluindo zeros à esquerda
                )
       
                email_pessoal = dados.get("E-mail pessoal", "")
       
                st.session_state["titulo_gerado"] = (
                    f"{nome_selecionado} __ "
                    f"{cpf_limpo} __ "
                    f"{email_pessoal} __ "
                    f"{titulo_doc}"
                )
        
            # ---------- TÍTULO GERADO ----------
            if "titulo_gerado" in st.session_state:
                st.markdown("#### 📄 Título gerado")
                st.code(st.session_state["titulo_gerado"])
        
        
        # ---------- BOTÃO QUE ABRE O MODAL (RESET TOTAL) ----------
        def abrir_modal_titulo():
            st.session_state["titulo_doc"] = ""
            st.session_state["nome_selecionado"] = ""
            st.session_state.pop("titulo_gerado", None)
            modal_titulo_doc()
        
        
        if st.button("📝 Título de doc para automação", use_container_width=True):
            abrir_modal_titulo()

        # --------------------------------------------------
        # AUTOMAÇÃO — DEMISSÃO POR COMUM ACORDO
        # --------------------------------------------------
        def substituir_texto_docx(doc, mapa):
            """
            Substitui chaves por valores em todo o documento,
            unindo runs para garantir que campos divididos sejam substituídos.
            """
        
            def substituir_em_paragrafo(paragrafo, mapa):
                # Junta todo o texto do parágrafo
                texto_completo = "".join(run.text for run in paragrafo.runs)
                for chave, valor in mapa.items():
                    if chave in texto_completo:
                        texto_completo = texto_completo.replace(chave, str(valor))
                # Remove runs antigas
                for run in paragrafo.runs:
                    run.text = ""
                # Adiciona texto atualizado como um único run
                paragrafo.add_run(texto_completo)
        
            # Parágrafos principais
            for p in doc.paragraphs:
                substituir_em_paragrafo(p, mapa)
        
            # Tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for p in celula.paragraphs:
                            substituir_em_paragrafo(p, mapa)
        
            # Cabeçalhos e rodapés
            for section in doc.sections:
                for p in section.header.paragraphs:
                    substituir_em_paragrafo(p, mapa)
                for p in section.footer.paragraphs:
                    substituir_em_paragrafo(p, mapa)

                
        # BOTÃO PRINCIPAL
        @st.dialog("📄 Demissão por comum acordo")  # Deixe em branco se não quiser título
        def modal_comum():

            st.markdown('<div class="modal_comum">', unsafe_allow_html=True)
            
            st.markdown("#### Preencha os dados abaixo")

            nome_selecionado = st.selectbox(
                "Nome do colaborador",
                sorted(df["Nome"].dropna().unique())
            )
     
            data_desligamento = st.date_input(
                "Data do desligamento",
                format="DD/MM/YYYY"
            )
        
            # BUSCA DADOS DA PESSOA
            dados_pessoa = df[df["Nome"] == nome_selecionado].iloc[0]
        
            cargo = dados_pessoa["Cargo"]
        
            # BOTÕES DE AÇÃO
            if st.button("✅ Gerar doc"):
        
                from docx import Document
                from io import BytesIO
        
                # Abre modelo
                doc = Document("Demissão por comum acordo.docx")
        
                mapa_substituicao = {
                    "{nome_completo}": nome_selecionado,
                    "{cargo}": cargo,
                    "{data}": data_desligamento.strftime("%d/%m/%Y")
                }

                # ✅ SUBSTITUI TEXTO (CORRETO)
                substituir_texto_docx(doc, mapa_substituicao)
        
                # SALVA EM MEMÓRIA
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
        
                st.success("Documento gerado com sucesso ✅")
       
                st.download_button(
                    label="⬇️ Baixar documento",
                    data=buffer,
                    file_name=f"Demissão - {nome_selecionado}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
            if st.button("❌ Cancelar"):
                st.session_state.gerar_demissao_comum = False
                st.rerun()

            st.markdown('</div>', unsafe_allow_html=True)

        if st.button("📄 Demissão por comum acordo", use_container_width=True):
            modal_comum()
        
        @st.dialog("📄 Aviso Prévio Indenizado")
        def modal_aviso_previo_indenizado():
        
            st.markdown("#### Preencha os dados")
        
            lista_nomes = sorted(df["Nome"].dropna().unique())
        
            nome = st.selectbox(
                "Nome do investidor",
                ["Selecione..."] + lista_nomes
            )
        
            data_desligamento = st.date_input(
                "Data do desligamento",
                format="DD/MM/YYYY"
            )
        
            data_homologacao = st.date_input(
                "Data da homologação",
                format="DD/MM/YYYY"
            )
        
            if st.button("📄 Gerar documento", use_container_width=True):
        
                if nome == "Selecione...":
                    st.warning("Selecione o investidor.")
                    return
        
                mapa = {
                    "{nome_selecionado}": nome,
                    "{data_desligamento}": data_desligamento.strftime("%d/%m/%Y"),
                    "{data_homologacao}": data_homologacao.strftime("%d/%m/%Y"),
                }
        
                arquivo = gerar_docx_com_substituicoes(
                    "Aviso prévio Indenizado.docx",
                    mapa
                )
        
                st.success("Documento gerado com sucesso!")
        
                st.download_button(
                    label="⬇️ Baixar documento",
                    data=arquivo,
                    file_name=f"Aviso prévio Indenizado - {nome}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        
        if st.button("📄 Aviso Prévio Indenizado", use_container_width=True):
            modal_aviso_previo_indenizado()

        def substituir_runs_paragrafos(doc, mapa):
            for p in doc.paragraphs:
                for run in p.runs:
                    for chave, valor in mapa.items():
                        if chave in run.text:
                            run.text = run.text.replace(chave, str(valor))
        
        
        def substituir_runs_tabelas(doc, mapa):
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                for chave, valor in mapa.items():
                                    if chave in run.text:
                                        run.text = run.text.replace(chave, str(valor))

        @st.dialog("🚌 Atualização do Vale Transporte")
        def modal_vale_transporte(df_pessoas):
        
            # =====================
            # INVESTIDOR
            # =====================
            nome_sel = st.selectbox(
                "Investidor",
                df_pessoas["Nome"].tolist()
            )
        
            cpf_sel = df_pessoas.loc[
                df_pessoas["Nome"] == nome_sel, "CPF"
            ].values[0]
        
            # =====================
            # ENDEREÇO
            # =====================
            cep = st.text_input("CEP")
            endereco = st.text_input("Endereço")
            numero = st.text_input("Número")
            bairro = st.text_input("Bairro")
            cidade = st.text_input("Cidade")
            uf = st.text_input("UF")
        
            # =====================
            # IDA
            # =====================
            st.divider()
            st.subheader("Residência → Trabalho")
        
            qtd_res = st.selectbox("Quantidade de transportes", [1,2,3,4], key="qtd_res")
        
            transportes_res = []
        
            for i in range(qtd_res):
                c1, c2, c3, c4 = st.columns(4)
        
                tipo = c1.selectbox(
                    "Tipo", ["Ônibus", "Metrô", "Trem"], key=f"tipo_res_{i}"
                )
                linha = c2.text_input("Linha", key=f"linha_res_{i}")
                valor = c3.number_input(
                    "Valor", min_value=0.0, step=0.01, key=f"valor_res_{i}"
                )
                inte = c4.number_input(
                    "Integração", min_value=0.0, step=0.01, key=f"inte_res_{i}"
                )
        
                transportes_res.append((tipo, linha, valor, inte))
        
            soma_linhas = len(transportes_res)
            soma_valor = sum(v for _,_,v,_ in transportes_res)
            soma_inte = sum(i for _,_,_,i in transportes_res)
        
            # =====================
            # VOLTA
            # =====================
            st.divider()
            st.subheader("Trabalho → Residência")
        
            qtd_tra = st.selectbox("Quantidade de transportes", [1,2,3,4], key="qtd_tra")
        
            transportes_tra = []
        
            for i in range(qtd_tra):
                c1, c2, c3, c4 = st.columns(4)
        
                tipo = c1.selectbox(
                    "Tipo", ["Ônibus", "Metrô", "Trem"], key=f"tipo_tra_{i}"
                )
                linha = c2.text_input("Linha", key=f"linha_tra_{i}")
                valor = c3.number_input(
                    "Valor", min_value=0.0, step=0.01, key=f"valor_tra_{i}"
                )
                inte = c4.number_input(
                    "Integração", min_value=0.0, step=0.01, key=f"inte_tra_{i}"
                )
        
                transportes_tra.append((tipo, linha, valor, inte))
        
            soma_linhas_tra = len(transportes_tra)
            soma_valor_tra = sum(v for _,_,v,_ in transportes_tra)
            soma_inte_tra = sum(i for _,_,_,i in transportes_tra)
        
            # =====================
            # TOTAIS
            # =====================
            soma_unit = soma_valor + soma_valor_tra
            soma_integracao = soma_inte + soma_inte_tra
        
            # =====================
            # DATA
            # =====================
            MESES_PT = {
                1:"janeiro",2:"fevereiro",3:"março",4:"abril",
                5:"maio",6:"junho",7:"julho",8:"agosto",
                9:"setembro",10:"outubro",11:"novembro",12:"dezembro"
            }
        
            hoje = date.today()
            data_extenso = f"{hoje.day} de {MESES_PT[hoje.month]} de {hoje.year}"

            import os
            from docx import Document
            
            CAMINHO_MODELO = os.path.join(
                os.getcwd(),
                "declaracao_vale_transporte_clt.docx"
            )
            
            # =====================
            # GERAR DOCUMENTO
            # =====================
            st.divider()
            c1, c2, c3 = st.columns([1, 2, 1])
            
            with c2:
                gerar = st.button("📄 Gerar documento", use_container_width=True)
            
            if gerar:
            
                mapa = {
                    "{nome}": nome_sel,
                    "{cpf}": cpf_sel,
                    "{cep}": cep,
                    "{endereço}": endereco,
                    "{número}": numero,
                    "{bairro}": bairro,
                    "{cidade}": cidade,
                    "{uf_estado}": uf,
                    "{soma_linhas}": str(soma_linhas),
                    "{soma_valor}": f"{soma_valor:.2f}",
                    "{soma_inte}": f"{soma_inte:.2f}",
                    "{soma_linhas_tra}": str(soma_linhas_tra),
                    "{soma_valor_tra}": f"{soma_valor_tra:.2f}",
                    "{soma_inte_tra}": f"{soma_inte_tra:.2f}",
                    "{soma_unit}": f"{soma_unit:.2f}",
                    "{soma_integracao}": f"{soma_integracao:.2f}",
                    "{data}": data_extenso
                }
            
                # 🔹 GARANTE CAMPOS EM BRANCO (IDA)
                for i in range(1, 5):
                    mapa.setdefault(f"{{transporte_{i}_res}}", "")
                    mapa.setdefault(f"{{linha_{i}_res}}", "")
                    mapa.setdefault(f"{{valor_{i}_res}}", "")
                    mapa.setdefault(f"{{inte_{i}_res}}", "")
            
                # 🔹 GARANTE CAMPOS EM BRANCO (VOLTA)
                for i in range(1, 5):
                    mapa.setdefault(f"{{transporte_{i}_tra}}", "")
                    mapa.setdefault(f"{{linha_{i}_tra}}", "")
                    mapa.setdefault(f"{{valor_{i}_tra}}", "")
                    mapa.setdefault(f"{{inte_{i}_tra}}", "")
            
                # 🔹 SOBRESCREVE IDA
                for i, (t, l, v, it) in enumerate(transportes_res, start=1):
                    mapa[f"{{transporte_{i}_res}}"] = t
                    mapa[f"{{linha_{i}_res}}"] = l
                    mapa[f"{{valor_{i}_res}}"] = f"{v:.2f}"
                    mapa[f"{{inte_{i}_res}}"] = f"{it:.2f}"
            
                # 🔹 SOBRESCREVE VOLTA
                for i, (t, l, v, it) in enumerate(transportes_tra, start=1):
                    mapa[f"{{transporte_{i}_tra}}"] = t
                    mapa[f"{{linha_{i}_tra}}"] = l
                    mapa[f"{{valor_{i}_tra}}"] = f"{v:.2f}"
                    mapa[f"{{inte_{i}_tra}}"] = f"{it:.2f}"
            
                doc = Document(CAMINHO_MODELO)

                substituir_runs_paragrafos(doc, mapa)
                substituir_runs_tabelas(doc, mapa)
                substituir_runs_header_footer(doc, mapa)
            
                nome_arquivo = f"Declaração de Vale Transporte CLT - {nome_sel}.docx"
                doc.save(nome_arquivo)
            
                with open(nome_arquivo, "rb") as f:
                    c1, c2, c3 = st.columns([1, 2, 1])
                    with c2:
                        st.download_button(
                            "⬇️ Download do documento",
                            f,
                            file_name=nome_arquivo,
                            use_container_width=True
                        )
                        
        if st.button("🚌 Atualização do Vale Transporte", use_container_width=True):
            modal_vale_transporte(df_pessoas=df)

