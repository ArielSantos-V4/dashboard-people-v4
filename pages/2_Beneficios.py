import streamlit as st
import pandas as pd
import plotly.express as px
import bcrypt
import altair as alt
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from docx import Document
from datetime import date

if "investidor_selecionado" not in st.session_state:
    st.session_state.investidor_selecionado = ""

def limpar_investidor():
    st.session_state.investidor_selecionado = ""

def formatar_cpf(valor):
    v = str(valor).replace(".0", "").zfill(11)
    if len(v) != 11:
        return ""
    return f"{v[:3]}.{v[3:6]}.{v[6:9]}-{v[9:]}"


def formatar_cnpj(valor):
    v = str(valor).replace(".0", "").zfill(14)
    if len(v) != 14:
        return ""
    return f"{v[:2]}.{v[2:5]}.{v[5:8]}/{v[8:12]}-{v[12:]}"

def render_table(df, *, dataframe=True, **kwargs):
    """
    Renderiza tabelas no Streamlit sem mostrar NaN / NaT / None,
    preservando os tipos originais do dataframe.
    """
    df_view = df.copy()

    # Substitui apenas para exibição
    df_view = df_view.where(pd.notna(df_view), "")

    if dataframe:
        st.dataframe(df_view, **kwargs)
    else:
        st.table(df_view)

def parse_data_br(coluna):
    return pd.to_datetime(coluna, dayfirst=True, errors="coerce")

from dateutil.relativedelta import relativedelta

def calcular_tempo_casa(data_inicio):
    if pd.isna(data_inicio):
        return ""

    hoje = pd.Timestamp.today().normalize()
    diff = relativedelta(hoje, data_inicio)

    return f"{diff.years} anos, {diff.months} meses e {diff.days} dias"

import unicodedata

def email_para_nome_arquivo(email):
    if not email:
        return ""

    email = unicodedata.normalize("NFKC", email)

    return (
        email
        .strip()
        .lower()
        .replace(" ", "")
    )

import re

def normalizar_cpf(cpf):
    if not cpf:
        return ""

    # remove tudo que não for número
    cpf = re.sub(r"\D", "", str(cpf))

    # garante 11 dígitos com zero à esquerda
    return cpf.zfill(11)

def gerar_hash_senha(senha):
    return bcrypt.hashpw(
        senha.encode("utf-8"),
        bcrypt.gensalt()
    ).decode("utf-8")

def gerar_alertas_investidor(linha):
    alertas = []

    # --- data de hoje (sem hora) ---
    hoje = pd.Timestamp.today().normalize()

    # --- status do plano ---
    status = str(linha["Situação no plano"]).strip()

    # =========================================================
    # ALERTA 1 — SOLICITAR DOCUMENTAÇÃO
    # status = Pendente → usa coluna "Solicitar documentação"
    # =========================================================
    data_solicitar = pd.to_datetime(
        linha["Solicitar documentação"],
        errors="coerce"
    )

    if status == "Pendente" and pd.notna(data_solicitar):
        dias = (data_solicitar - hoje).days

        if dias < 0:
            alertas.append((
                "error",
                "Plano de saúde e dental 🤕\n"
                "Solicitação de documentação em atraso. Verificar com urgência!"
            ))
        elif dias == 0:
            alertas.append((
                "warning",
                "Plano de saúde e dental ❤️‍🩹\n"
                "Hoje é a data limite para solicitar a documentação!"
            ))
        elif dias <= 15:
            alertas.append((
                "info",
                f"Plano de saúde e dental ❤️‍🩹\n"
                f"Faltam {dias} dias para solicitar a documentação ao investidor"
            ))

    # =========================================================
    # ALERTA 2 — ENVIAR NO EB
    # status = Aguardando docs → usa coluna "Enviar no EB"
    # =========================================================
    data_enviar_eb = pd.to_datetime(
        linha["Enviar no EB"],
        errors="coerce"
    )

    if status == "Aguardando docs" and pd.notna(data_enviar_eb):
        dias = (data_enviar_eb - hoje).days

        if dias < 0:
            alertas.append((
                "error",
                "Plano de saúde e dental 🤕\n"
                "Envio à EB em atraso. Verificar com urgência!"
            ))
        elif dias == 0:
            alertas.append((
                "warning",
                "Plano de saúde e dental ❤️‍🩹\n"
                "Hoje é a data limite para enviar à EB"
            ))
        elif dias <= 15:
            alertas.append((
                "info",
                f"Plano de saúde e dental ❤️‍🩹\n"
                f"Faltam {dias} dias para enviar à EB"
            ))

    if status == "Aguardando DBL":
        alertas.append(("info",
            "Plano de saúde e dental quase prontos! 🤩"
            "Acompanhar movimentação no portal EB"
        ))
    
    # -------------------------
    # ALERTA — Aniversário
    # -------------------------
    nascimento_raw = linha.get("Data de nascimento", "")
    
    nascimento = pd.to_datetime(
        nascimento_raw,
        errors="coerce",
        dayfirst=True
    )
    
    if pd.notna(nascimento):
        nascimento = pd.Timestamp(nascimento).normalize()
    
        if nascimento.month == hoje.month:
            if nascimento.day == hoje.day:
                alertas.append((
                    "info",
                    "Lembrete de Aniversário! 🎉\n"
                    "HOJE é aniversário do investidor!!"
                ))
            else:
                alertas.append((
                    "info",
                    "Lembrete de Aniversário! 🎉\n"
                    "Este investidor faz aniversário neste mês"
                ))

    # -------------------------
    # ALERTA 3 — Contrato
    # -------------------------
    fim_contrato_raw = linha.get("Térm previsto", "")

    fim_contrato = pd.to_datetime(
        fim_contrato_raw,
        errors="coerce",
        dayfirst=True
    )
    
    if pd.notna(fim_contrato):
        fim_contrato = pd.Timestamp(fim_contrato).normalize()
        dias = (fim_contrato - hoje).days
    
    if pd.notna(fim_contrato):
        dias = (fim_contrato - hoje).days

        if dias < 0:
            alertas.append(("error",
                "Contrato vencido! 🚨"
                "Verificar com urgência!"
            ))
        elif dias <= 30:
            alertas.append(("warning",
                f"Alerta! ⚠️"
                f"O contrato se encerra em {dias} dia(s)."
            ))

    # -------------------------
    # ALERTA 4 — MEI
    # -------------------------
    if linha.get("Modalidade PJ", "") == "MEI":
        alertas.append(("warning",
            "Atenção! Investidor ainda se encontra na modalidade MEI 😬"
        ))

    return alertas

st.markdown("""
<style>
/* Modal específico da consulta individual */
div[role="dialog"]:has(.modal-investidor) {
    width: 95vw !important;
    max-width: 95vw !important;
}

/* Altura maior (opcional) */
div[role="dialog"]:has(.modal-investidor) > div {
    max-height: 90vh !important;
}
    
</style>
""", unsafe_allow_html=True)

from docx import Document
from io import BytesIO

def gerar_docx_com_substituicoes(caminho_modelo, substituicoes):
    doc = Document(caminho_modelo)

    for paragrafo in doc.paragraphs:
        for run in paragrafo.runs:
            for chave, valor in substituicoes.items():
                if chave in run.text:
                    run.text = run.text.replace(chave, valor)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for run in paragrafo.runs:
                        for chave, valor in substituicoes.items():
                            if chave in run.text:
                                run.text = run.text.replace(chave, valor)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

from docx import Document

def gerar_vale_transporte(dados):
    # Abre o documento ORIGINAL (modelo)
    doc = Document("modelo_vale_transporte.docx")

    for p in doc.paragraphs:
        if "{{NOME}}" in p.text:
            p.text = p.text.replace("{{NOME}}", dados["nome"])

        if "{{CPF}}" in p.text:
            p.text = p.text.replace("{{CPF}}", dados["cpf"])

        if "{{VALOR}}" in p.text:
            p.text = p.text.replace("{{VALOR}}", dados["valor"])

    doc.save("vale_transporte_final.docx")

def substituir_runs_header_footer(doc, mapa):
    for section in doc.sections:
        # HEADER
        for p in section.header.paragraphs:
            for run in p.runs:
                for chave, valor in mapa.items():
                    if chave in run.text:
                        run.text = run.text.replace(chave, str(valor))

        # FOOTER
        for p in section.footer.paragraphs:
            for run in p.runs:
                for chave, valor in mapa.items():
                    if chave in run.text:
                        run.text = run.text.replace(chave, str(valor))
                        
def render():
    
    # --------------------------------------------------
    # ABAS
    # --------------------------------------------------
    aba_benefícios = st.tabs([
        "🎁 Benefícios"
    ])
    
    # --------------------------------------------------
    # ABA BENEFICIOS
    # --------------------------------------------------
    
    with aba_benefícios:
    
        st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    
        # --------------------------------------------------
        # TOPO
        # --------------------------------------------------
        col_logo, col_title = st.columns([1, 6])
    
        with col_logo:
            st.image("LOGO VERMELHO.png", width=120)
    
        with col_title:
            st.markdown(
                "<h1> Gestão de Benefícios </h1>"
                "<h3 style='color:#ccc;'>V4 Company</h3>",
                unsafe_allow_html=True
            )
    
        st.markdown("---")
    
        from datetime import datetime, timedelta
        import altair as alt
    
        # --------------------------------------------------
        # LAYOUT — BENEFÍCIOS
        # --------------------------------------------------
        col_grafico, col_consulta = st.columns([4, 6])
    
        # ---------------------------------
        # COLUNA 1 — GRÁFICO SITUAÇÃO NO PLANO
        # ---------------------------------
        with col_grafico:
    
            st.markdown("<h3 style='margin-bottom:20px'>📊 Status no plano</h3>", unsafe_allow_html=True)
            st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
        
            if "Situação no plano" in df.columns:
        
                df_plano = (
                    df["Situação no plano"]
                    .fillna("Não informado")
                    .value_counts()
                    .reset_index()
                )
        
                df_plano.columns = ["Situação", "Quantidade"]
                total = df_plano["Quantidade"].sum()
                df_plano["Percentual"] = (df_plano["Quantidade"] / total) * 100
        
                grafico_plano = (
                    alt.Chart(df_plano)
                    .mark_arc(innerRadius=80, outerRadius=130, stroke=None)
                    .encode(
                        theta="Quantidade:Q",
                        color=alt.Color(
                            "Situação:N",
                            scale=alt.Scale(
                                range=[
                                    "#2E8B57",
                                    "#FFA500",
                                    "#8A2BE2",
                                    "#DC143C",
                                    "#8B4513",
                                    "#808080",
                                ]
                            ),
                            legend=alt.Legend(
                                title="Situação",
                                orient="bottom",
                                columns=2,
                                offset=20
                            )
                        ),
                        tooltip=[
                            alt.Tooltip("Situação:N", title="Situação"),
                            alt.Tooltip("Quantidade:Q", title="Qtd"),
                            alt.Tooltip("Percentual:Q", title="%", format=".1f"),
                        ],
                    )
                    .properties(width=320, height=380)
                )
        
                st.altair_chart(grafico_plano, use_container_width=True)
        
            else:
                st.warning("Coluna 'Situação no plano' não encontrada.")
    
        # ---------------------------------
        # COLUNA 2 — CONSULTA CARTEIRINHAS
        # ---------------------------------
        with col_consulta:
        
            st.markdown("### 🔎 Consulta de carteirinhas")
        
            nome_beneficio = st.selectbox(
                "Selecione o investidor",
                options=[""] + sorted(df["Nome"].dropna().unique()),
                placeholder="Digite ou selecione um nome"
            )
        
            consultar = st.button("Consultar carteirinhas", use_container_width=True)
        
            if nome_beneficio and consultar:
        
                dados = df[df["Nome"] == nome_beneficio].iloc[0]
        
                cart_med = str(dados.get("Carteirinha médico", "")).strip()
                oper_med = str(dados.get("Operadora Médico", "")).strip()
                cart_odo = str(dados.get("Carteirinha odonto", "")).strip()
                oper_odo = str(dados.get("Operadora Odonto", "")).strip()
                situacao = str(dados.get("Situação no plano", "Não informado"))
        
                # 🔴 CASO NÃO TENHA CARTEIRINHA (NÃO ATIVO)
                if not cart_med and not cart_odo:
        
                    st.markdown(
                        f"""
                        <div style="
                            position: relative;
                            padding: 25px;
                            border-radius: 12px;
                            background: rgba(0,0,0,0.55);
                            backdrop-filter: blur(6px);
                            -webkit-backdrop-filter: blur(6px);
                            color: white;
                            text-align: center;
                        ">
                            <h4>⚠️ Investidor não ativo no plano</h4>
                            <p>Este investidor não possui carteirinhas ativas no momento.</p>
                            <hr style="opacity:0.2;">
                            <div style="
                                margin-top: 12px;
                                padding: 10px;
                                border-radius: 8px;
                                background-color: #8B0000;
                                color: white;
                                font-weight: bold;
                            ">
                                Situação atual no plano: {situacao}
                            </div>
    
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
        
                # 🟢 CASO TENHA CARTEIRINHA
                else:
                    st.text_input(
                        "Carteirinha médico",
                        cart_med if cart_med else "—",
                        disabled=True
                    )
                    st.text_input(
                        "Operadora médico",
                        oper_med if oper_med else "—",
                        disabled=True
                    )
        
                    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
        
                    st.text_input(
                        "Carteirinha odonto",
                        cart_odo if cart_odo else "—",
                        disabled=True
                    )
                    st.text_input(
                        "Operadora odonto",
                        oper_odo if oper_odo else "—",
                        disabled=True
                    )
    
        st.markdown("---")
    
        # ==============================
        # BLOCO — RELATÓRIOS & AÇÕES
        # ==============================
            
        # -------- GRID PRINCIPAL --------
        col_relatorios, col_acoes = st.columns([7, 3])
        
        # ==============================
        # COLUNA ESQUERDA — RELATÓRIOS
        # ==============================
        with col_relatorios:
            st.markdown("### 📊 Relatórios")
        
            abas = st.tabs([
                "⏰ Pendentes",
                "📂 Aguardando docs",
                "📩 Enviar para DBL",
                "🆗 Aguardando ativação"
            ])
        
            with abas[0]:
                st.markdown("#### Investidores com documentação pendente")
            
                # --- FILTRO: somente pendentes ---
                df_pendentes = df[
                    (df["Situação no plano"] == "Pendente") &
                    (df["Modalidade PJ"] != "MEI")
                ]
            
                # --- SELEÇÃO DAS COLUNAS ---
                tabela_docs = df_pendentes[[
                    "Nome",
                    "E-mail corporativo",
                    "Modelo de contrato",
                    "Solicitar documentação"
                ]]
            
                st.dataframe(
                    tabela_docs,
                    use_container_width=True,
                    hide_index=True
                )
    
            with abas[1]:
                st.markdown("#### Aguardando envio da documentação")
            
                # --- FILTRO: somente pendentes ---
                df_pendentes = df[df["Situação no plano"] == "Aguardando docs"]
            
                # --- SELEÇÃO DAS COLUNAS ---
                tabela_docs = df_pendentes[[
                    "Nome",
                    "E-mail corporativo",
                    "Modelo de contrato",
                    "Enviar no EB"
                ]]
            
                st.dataframe(
                    tabela_docs,
                    use_container_width=True,
                    hide_index=True
                )
                
            with abas[2]:
                st.markdown("#### Investidores para envio à DBL")
            
                # --- FILTRO: aguardando documentação ---
                df_dbl = df[df["Situação no plano"] == "Enviar à DBL"]
            
                # --- SELEÇÃO DAS COLUNAS ---
                tabela_dbl = df_dbl[[
                    "Nome",
                    "E-mail corporativo",
                    "Modelo de contrato",
                    "Enviar no EB"
                ]]
            
                st.dataframe(
                    tabela_dbl,
                    use_container_width=True,
                    hide_index=True
                )
    
        
            with abas[3]:
                st.markdown("#### Investidores aguardando retorno da DBL")
            
                # --- FILTRO: aguardando DBL ---
                df_dbl_status = df[df["Situação no plano"] == "Aguardando DBL"]
            
                # --- COLUNAS EXIBIDAS ---
                tabela_dbl_status = df_dbl_status[[
                    "Nome",
                    "E-mail corporativo",
                    "Modelo de contrato"
                ]]
            
                st.dataframe(
                    tabela_dbl_status,
                    use_container_width=True,
                    hide_index=True
                )
    
    
        with col_acoes:
            # ==============================
            # AÇÃO — INCLUSÃO SUBFATURA
            # ==============================
            
            from docx import Document
            import re
            from datetime import datetime, date
            
            MESES_PT = {
                1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
                5: "maio", 6: "junho", 7: "julho", 8: "agosto",
                9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
            }
            
            def substituir_texto(paragraphs, mapa):
                for p in paragraphs:
                    for run in p.runs:
                        for chave, valor in mapa.items():
                            if chave in run.text:
                                run.text = run.text.replace(chave, str(valor))
    
            def formatar_cnpj(cnpj):
                # Converte para string e remove .0 se vier como float
                cnpj_str = str(cnpj).replace(".0", "")
                
                # Remove tudo que não for número
                cnpj_numeros = re.sub(r"\D", "", cnpj_str)
            
                # Garante 14 dígitos (com zeros à esquerda se necessário)
                cnpj_numeros = cnpj_numeros.zfill(14)
            
                return (
                    f"{cnpj_numeros[0:2]}."
                    f"{cnpj_numeros[2:5]}."
                    f"{cnpj_numeros[5:8]}/"
                    f"{cnpj_numeros[8:12]}-"
                    f"{cnpj_numeros[12:14]}"
                )
                
            # -------- BOTÃO PRINCIPAL --------
            st.markdown("### ⚙️ Ações")
    
            @st.dialog("📄 Gerar Inclusão Subfatura")
            def modal_inclusao_subfatura():
            
                nomes = sorted(df["Nome"].dropna().unique())
                nome_escolhido = st.selectbox("Selecione o investidor", nomes)
            
                data_vigencia = st.date_input(
                    "Data de início da vigência",
                    format="DD/MM/YYYY"
                )
            
                st.markdown("<br>", unsafe_allow_html=True)
            
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    gerar = st.button("✅ Gerar", use_container_width=True)
            
                if gerar:
            
                    dados = df[df["Nome"] == nome_escolhido].iloc[0]
            
                    razao_social = str(dados["Razão social"])
                    cnpj = formatar_cnpj(dados["CNPJ"])
                    cpf = normalizar_cpf(dados["CPF"])
                    email_pessoal = str(dados["E-mail pessoal"])
                    email_arquivo = email_para_nome_arquivo(email_pessoal)
                    modelo_contrato = str(dados["Modelo de contrato"])
            
                    # -------- VALIDAÇÃO PJ --------
                    if "PJ" not in modelo_contrato.upper():
                        st.warning(
                            f"⚠️ **{nome_escolhido}** não possui contrato PJ.\n\n"
                            f"Modelo atual: **{modelo_contrato}**"
                        )
            
                    # -------- ABRE TEMPLATE --------
                    doc = Document("Subfatura.docx")
            
                    vigencia_formatada = data_vigencia.strftime("%d/%m/%Y")
            
                    hoje = date.today()
                    data_assinatura = f"{hoje.day} de {MESES_PT[hoje.month]} de {hoje.year}"
            
                    mapa = {
                        "{RAZAO_SOCIAL}": razao_social,
                        "{CNPJ}": cnpj,
                        "{VIGENCIA}": vigencia_formatada,
                        "{DATA}": data_assinatura
                    }
            
                    substituir_texto(doc.paragraphs, mapa)
            
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                substituir_texto(cell.paragraphs, mapa)
            
                    for section in doc.sections:
                        substituir_texto(section.header.paragraphs, mapa)
            
                    cpf_limpo = re.sub(r"\D", "", cpf)
            
                    nome_arquivo = (
                        f"{nome_escolhido} __ {cpf_limpo} __ {email_arquivo} __ Inclusão Subfatura.docx"
                    )
            
                    doc.save(nome_arquivo)
            
                    col_btn1, col_btn2 = st.columns(2)
    
                    with col_btn1:
                        with open(nome_arquivo, "rb") as f:
                            st.download_button(
                                "⬇️ Download",
                                data=f,
                                file_name=nome_arquivo,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    
                    with col_btn2:
                        st.link_button(
                            "🔁 Converter PDF",
                            "https://www.ilovepdf.com/pt/word_para_pdf",
                            use_container_width=True
                        )
    
            
                    st.success("Inclusão Subfatura gerada com sucesso ✅")
                    
            if st.button("📄 Gerar Inclusão Subfatura", use_container_width=True):
                modal_inclusao_subfatura()
           
            # ==============================
            # AÇÃO — GERAR TERMO DE SUBESTIPULANTE
            # ==============================
            @st.dialog("📄 Gerar Termo de Subestipulante")
            def modal_subestipulante():
    
                st.markdown('<div class="modal_subestipulante">', unsafe_allow_html=True)
           
                nomes = sorted(df["Nome"].dropna().unique())
                nome_escolhido = st.selectbox(
                    "Selecione o investidor",
                    nomes,
                    key="nome_termo"
                )
            
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    gerar_termo = st.button(
                        "✅ Gerar Termo",
                        use_container_width=True,
                        key="btn_gerar_termo"
                    )
            
                if gerar_termo:
            
                    dados = df[df["Nome"] == nome_escolhido].iloc[0]
            
                    razao_social = str(dados["Razão social"])
                    cnpj = formatar_cnpj(dados["CNPJ"])
                    cpf = normalizar_cpf(dados["CPF"])
                    email_pessoal = str(dados["E-mail pessoal"])
                    email_arquivo = email_para_nome_arquivo(email_pessoal)
            
                    # -------- ABRE TEMPLATE --------
                    doc = Document("Termo de integração de subestipulante.docx")
            
                    hoje = date.today()
                    data_assinatura = f"{hoje.day} de {MESES_PT[hoje.month]} de {hoje.year}"
            
                    mapa = {
                        "{RAZAO_SOCIAL}": razao_social,
                        "{CNPJ}": cnpj,
                        "{DATA}": data_assinatura
                    }
            
                    # Parágrafos normais
                    substituir_texto(doc.paragraphs, mapa)
                    
                    # Tabelas
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                substituir_texto(cell.paragraphs, mapa)
                    
                    # Cabeçalho
                    for section in doc.sections:
                        substituir_texto(section.header.paragraphs, mapa)
    
            
                    cpf_limpo = re.sub(r"\D", "", cpf)
            
                    nome_arquivo = (
                        f"{nome_escolhido} __ {cpf_limpo} __ {email_arquivo} __ Termo Subestipulante.docx"
                    )
            
                    doc.save(nome_arquivo)
            
                    col_btn1, col_btn2 = st.columns(2)
            
                    with col_btn1:
                        with open(nome_arquivo, "rb") as f:
                            st.download_button(
                                "⬇️ Download",
                                data=f,
                                file_name=nome_arquivo,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
            
                    with col_btn2:
                        st.link_button(
                            "🔁 Converter PDF",
                            "https://www.ilovepdf.com/pt/word_para_pdf",
                            use_container_width=True
                        )
            
                    st.success("Termo de Subestipulante gerado com sucesso ✅")
                st.markdown('</div>', unsafe_allow_html=True)
    
            if st.button("📄 Gerar Termo de Subestipulante", use_container_width=True):
                modal_subestipulante()
    
            # ==============================
            # AÇÃO — GERAR TERMO DE NÃO ADESÃO
            # ==============================
            @st.dialog("📄 Gerar Termo de Não Adesão")
            def modal_nao_adesao():
                
                st.markdown('<div class="modal-nao-adesao">', unsafe_allow_html=True) 
                    
                nomes = sorted(df["Nome"].dropna().unique())
                nome_escolhido = st.selectbox(
                    "Selecione o investidor",
                    nomes,
                    key="nome_nao_adesao"
                )
            
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    gerar_nao_adesao = st.button(
                        "✅ Gerar Termo",
                        use_container_width=True,
                        key="btn_gerar_nao_adesao"
                    )
            
                if gerar_nao_adesao:
            
                    dados = df[df["Nome"] == nome_escolhido].iloc[0]
            
                    razao_social = str(dados["Razão social"])
                    cnpj = formatar_cnpj(dados["CNPJ"])
            
                    hoje = date.today()
                    data_assinatura = f"{hoje.day} de {MESES_PT[hoje.month]} de {hoje.year}"
            
                    mapa = {
                        "{RAZAO_SOCIAL}": razao_social,
                        "{CNPJ}": cnpj,
                        "{DATA}": data_assinatura
                    }
    
                    doc = Document("Termo de não adesão - Plano de Saúde e Dental.docx")
            
                    # Corpo
                    substituir_texto(doc.paragraphs, mapa)
            
                    # Tabelas (segurança extra)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                substituir_texto(cell.paragraphs, mapa)
            
                    # Cabeçalho e rodapé
                    for section in doc.sections:
                        substituir_texto(section.header.paragraphs, mapa)
                        substituir_texto(section.footer.paragraphs, mapa)
            
                    nome_arquivo = f"Termo de não adesão ao plano - {nome_escolhido}.docx"
            
                    doc.save(nome_arquivo)
            
                    col_btn1, col_btn2 = st.columns(2)
            
                    with col_btn1:
                        with open(nome_arquivo, "rb") as f:
                            st.download_button(
                                "⬇️ Download",
                                data=f,
                                file_name=nome_arquivo,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
            
                    with col_btn2:
                        st.link_button(
                            "🔁 Converter PDF",
                            "https://www.ilovepdf.com/pt/word_para_pdf",
                            use_container_width=True
                        )
            
                    st.success("Termo de Não Adesão gerado com sucesso ✅")
                    
                st.markdown('</div>', unsafe_allow_html=True)
            
            if st.button("📄 Gerar Termo de Não Adesão", use_container_width=True):
                modal_nao_adesao() 
    
            # ==============================
            # AÇÃO — EXCLUSÃO SUBFATURA
            # ==============================
            
            import streamlit as st
            import pandas as pd
            import re
            from datetime import date
            from docx import Document
            import gspread
            from google.oauth2.service_account import Credentials
            
            # ------------------------------
            # CONFIG GOOGLE SHEETS
            # ------------------------------
            
            def carregar_desligados_google_sheets():
            
                scopes = [
                    "https://www.googleapis.com/auth/spreadsheets",
                    "https://www.googleapis.com/auth/drive"
                ]
            
                creds = Credentials.from_service_account_file(
                    "credenciais_google.json",  # <-- ajuste aqui
                    scopes=scopes
                )
            
                client = gspread.authorize(creds)
            
                spreadsheet = client.open_by_key(
                    "ID_DA_PLANILHA"  # <-- ajuste aqui
                )
            
                worksheet = spreadsheet.get_worksheet_by_id(1422602176)
            
                dados = worksheet.get_all_records()
                return pd.DataFrame(dados)
            
            
            # ------------------------------
            # AUXILIARES
            # ------------------------------
            
            MESES_PT = {
                1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
                5: "maio", 6: "junho", 7: "julho", 8: "agosto",
                9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
            }
            
            def substituir_texto(paragraphs, mapa):
                for p in paragraphs:
                    for run in p.runs:
                        for chave, valor in mapa.items():
                            if chave in run.text:
                                run.text = run.text.replace(chave, str(valor))
            
            def formatar_cnpj(cnpj):
                cnpj_str = str(cnpj).replace(".0", "")
                cnpj_numeros = re.sub(r"\D", "", cnpj_str).zfill(14)
            
                return (
                    f"{cnpj_numeros[0:2]}."
                    f"{cnpj_numeros[2:5]}."
                    f"{cnpj_numeros[5:8]}/"
                    f"{cnpj_numeros[8:12]}-"
                    f"{cnpj_numeros[12:14]}"
                )
            
            def normalizar_cpf(cpf):
                cpf_str = str(cpf).replace(".0", "")
                return re.sub(r"\D", "", cpf_str).zfill(11)
            
            def email_para_nome_arquivo(email):
                return email.replace("@", "_").replace(".", "_").lower()
            
            
            # ------------------------------
            # UI
            # ------------------------------
            
            @st.dialog("📄 Gerar Exclusão Subfatura")
            def modal_exclusao_subfatura():
            
                df_desligados = carregar_desligados_google_sheets()
            
                nomes = sorted(df_desligados["Nome"].dropna().unique())
                nome_escolhido = st.selectbox("Selecione o investidor", nomes)
            
                data_exclusao = st.date_input(
                    "Data de exclusão",
                    format="DD/MM/YYYY"
                )
            
                st.markdown("<br>", unsafe_allow_html=True)
            
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    gerar = st.button("✅ Gerar", use_container_width=True)
            
                if gerar:
            
                    dados = df_desligados[df_desligados["Nome"] == nome_escolhido].iloc[0]
            
                    razao_social = str(dados["Razão social"])
                    cnpj = formatar_cnpj(dados["CNPJ"])
                    cpf = normalizar_cpf(dados["CPF"])
                    email_pessoal = str(dados["E-mail pessoal"])
                    email_arquivo = email_para_nome_arquivo(email_pessoal)
                    modelo_contrato = str(dados["Modelo de contrato"])
            
                    # -------- VALIDAÇÃO PJ --------
                    if "PJ" not in modelo_contrato.upper():
                        st.warning(
                            f"⚠️ **{nome_escolhido}** não possui contrato PJ.\n\n"
                            f"Modelo atual: **{modelo_contrato}**"
                        )
            
                    # -------- TEMPLATE --------
                    doc = Document("Exclusao_Subfatura.docx")
            
                    data_exclusao_formatada = data_exclusao.strftime("%d/%m/%Y")
            
                    hoje = date.today()
                    data_assinatura = f"{hoje.day} de {MESES_PT[hoje.month]} de {hoje.year}"
            
                    mapa = {
                        "{RAZAO_SOCIAL}": razao_social,
                        "{CNPJ}": cnpj,
                        "{DATA_EXCLUSAO}": data_exclusao_formatada,
                        "{DATA}": data_assinatura
                    }
            
                    substituir_texto(doc.paragraphs, mapa)
            
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                substituir_texto(cell.paragraphs, mapa)
            
                    for section in doc.sections:
                        substituir_texto(section.header.paragraphs, mapa)
            
                    cpf_limpo = re.sub(r"\D", "", cpf)
            
                    nome_arquivo = (
                        f"{nome_escolhido} __ {cpf_limpo} __ {email_arquivo} __ Exclusão Subfatura.docx"
                    )
            
                    doc.save(nome_arquivo)
            
                    col_btn1, col_btn2 = st.columns(2)
            
                    with col_btn1:
                        with open(nome_arquivo, "rb") as f:
                            st.download_button(
                                "⬇️ Download",
                                data=f,
                                file_name=nome_arquivo,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
            
                    with col_btn2:
                        st.link_button(
                            "🔁 Converter PDF",
                            "https://www.ilovepdf.com/pt/word_para_pdf",
                            use_container_width=True
                        )
            
                    st.success("Exclusão Subfatura gerada com sucesso ✅")
            
            
            if st.button("📄 Gerar Exclusão Subfatura", use_container_width=True):
                modal_exclusao_subfatura()
    
